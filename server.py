#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Serveur unique pour le CRM Code-Barres
Sert à la fois l'API REST et les fichiers statiques (HTML/CSS/JS)
Base de données SQLite locale
"""

import sys
import io

# Forcer l'encodage UTF-8 pour stdout et stderr sur Windows
if sys.platform == 'win32':
    # Réencoder stdout et stderr en UTF-8
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    else:
        # Pour les versions Python plus anciennes
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from flask import Flask, request, jsonify, send_from_directory, Response, send_file
from flask_cors import CORS
import sqlite3
import os
import secrets
from datetime import datetime, timedelta
import requests
import threading
import queue
import json
import base64
import re
import urllib.parse
from io import BytesIO
from functools import wraps

# Charger les variables d'environnement depuis le fichier .env
try:
    from dotenv import load_dotenv
    load_dotenv()
    print('[CONFIG] Fichier .env chargé')
except ImportError:
    print('[CONFIG] python-dotenv non installé - variables d\'environnement système utilisées')

# ==================== CONFIGURATION VIA VARIABLES D'ENVIRONNEMENT ====================

# Mode de l'application : 'development' ou 'production'
APP_MODE = os.environ.get('APP_MODE', 'development')

# Origines CORS autorisées (séparées par des virgules)
CORS_ORIGINS = os.environ.get('CORS_ORIGINS', 'http://localhost:3000,http://localhost:3001,http://127.0.0.1:3000,http://127.0.0.1:3001')
CORS_ORIGINS_LIST = [origin.strip() for origin in CORS_ORIGINS.split(',') if origin.strip()]

# En mode développement, on peut autoriser toutes les origines
if APP_MODE == 'development':
    CORS_ORIGINS_LIST = ["*"]

# Port du serveur
SERVER_PORT = int(os.environ.get('SERVER_PORT', 5000))

print(f'[CONFIG] Mode: {APP_MODE}')
print(f'[CONFIG] CORS Origins: {CORS_ORIGINS_LIST}')
print(f'[CONFIG] Port: {SERVER_PORT}')

# Fonction d'affichage sécurisée pour Windows
def safe_print(msg):
    """Print qui ne plante jamais sur Windows (Errno 22)"""
    try:
        print(msg)
    except (OSError, IOError, UnicodeEncodeError):
        try:
            # Essayer avec ASCII uniquement
            print(msg.encode('ascii', errors='replace').decode('ascii'))
        except:
            pass  # Abandonner silencieusement

def safe_traceback():
    """Afficher le traceback de manière sécurisée sur Windows"""
    try:
        import traceback
        import io
        buffer = io.StringIO()
        traceback.print_exc(file=buffer)
        safe_print(buffer.getvalue())
    except:
        pass  # Ignorer les erreurs de traceback

# ==================== GESTION DES IMAGES ====================

# Utiliser un chemin absolu basé sur le répertoire du script (évite les problèmes Windows/OneDrive)
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(SCRIPT_DIR, 'data', 'images')

def ensure_images_dir():
    """Créer le dossier images s'il n'existe pas"""
    try:
        os.makedirs(IMAGES_DIR, exist_ok=True)
        return True
    except Exception as e:
        safe_print(f'[IMG] Erreur création dossier images: {e}')
        return False

def save_uploaded_file(file_storage, serial_number, index=0):
    """
    Sauvegarder un fichier uploadé directement sur disque.
    Simple et fiable - pas de Base64.
    """
    try:
        safe_print(f'[IMG] save_uploaded_file: debut')
        
        if not ensure_images_dir():
            safe_print(f'[IMG] Erreur: impossible de creer le dossier {IMAGES_DIR}')
            return None

        if not file_storage or not file_storage.filename:
            safe_print('[IMG] Erreur: file_storage invalide')
            return None

        # Extension depuis le nom de fichier original
        original_name = file_storage.filename
        ext = original_name.rsplit('.', 1)[-1].lower() if '.' in original_name else 'jpg'
        if ext not in ['jpg', 'jpeg', 'png', 'gif', 'webp']:
            ext = 'jpg'

        # Nom de fichier simple et unique
        unique_id = secrets.token_hex(6)
        safe_serial = re.sub(r'[^a-zA-Z0-9]', '', str(serial_number or 'img'))[:20]
        filename = f"{safe_serial}_{unique_id}.{ext}"
        filepath = os.path.join(IMAGES_DIR, filename)

        safe_print(f'[IMG] Sauvegarde vers: {filepath}')
        
        # Sauvegarder directement
        file_storage.save(filepath)
        safe_print(f'[IMG] Fichier sauvegarde OK: {filename}')
        return f"/api/images/{filename}"

    except Exception as e:
        safe_print(f'[IMG] Exception sauvegarde: {e}')
        import traceback
        traceback.print_exc()
        return None

def process_images_for_storage(image_data, serial_number):
    """
    Traiter les chemins d'images existants.
    Retourne un JSON array des chemins API ou None.
    """
    if not image_data:
        return None

    try:
        # Déjà un chemin API ou JSON de chemins
        if isinstance(image_data, str):
            if image_data.startswith('/api/images/'):
                return json.dumps([image_data])
            if image_data.startswith('['):
                try:
                    paths = json.loads(image_data)
                    if isinstance(paths, list):
                        valid = [p for p in paths if isinstance(p, str) and p.startswith('/api/images/')]
                        return json.dumps(valid) if valid else None
                except:
                    pass
        
        if isinstance(image_data, list):
            valid = [p for p in image_data if isinstance(p, str) and p.startswith('/api/images/')]
            return json.dumps(valid) if valid else None

        return None
    except Exception as e:
        safe_print(f'[IMG] Erreur traitement images: {e}')
        return None

# ==================== FONCTIONS DE VALIDATION ====================

def validate_required_fields(data, required_fields):
    """Valider que les champs requis sont présents et non vides"""
    missing = []
    for field in required_fields:
        if field not in data or data[field] is None or (isinstance(data[field], str) and not data[field].strip()):
            missing.append(field)
    return missing

def validate_email(email):
    """Valider le format d'un email"""
    if not email:
        return True  # Email optionnel
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def validate_phone(phone):
    """Valider le format d'un numéro de téléphone"""
    if not phone:
        return True  # Téléphone optionnel
    # Accepte les formats courants : +33, 06, 07, etc.
    pattern = r'^[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}$'
    return re.match(pattern, phone.replace(' ', '')) is not None

def sanitize_string(value, max_length=500):
    """Nettoyer et limiter la longueur d'une chaîne"""
    if value is None:
        return None
    if not isinstance(value, str):
        value = str(value)
    # Supprimer les caractères de contrôle dangereux
    value = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', value)
    # S'assurer que la chaîne peut être encodée en UTF-8
    try:
        # Tester l'encodage UTF-8
        value.encode('utf-8')
    except UnicodeEncodeError:
        # Si l'encodage échoue, remplacer les caractères problématiques
        value = value.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
    return value[:max_length].strip()

def validate_positive_number(value, allow_zero=True):
    """Valider qu'une valeur est un nombre positif"""
    try:
        num = float(value)
        if allow_zero:
            return num >= 0
        return num > 0
    except (TypeError, ValueError):
        return False

# ==================== RECHERCHE TESSERACT ====================

def find_tesseract():
    """Trouver le chemin de Tesseract selon l'OS"""
    # D'abord vérifier la variable d'environnement
    env_path = os.environ.get('TESSERACT_PATH')
    if env_path and os.path.exists(env_path):
        return env_path
    
    # Chemins par défaut selon l'OS
    if os.name == 'nt':  # Windows
        default_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            os.path.expanduser(r'~\AppData\Local\Tesseract-OCR\tesseract.exe'),
        ]
    else:  # Linux/Mac
        default_paths = [
            '/usr/bin/tesseract',
            '/usr/local/bin/tesseract',
            '/opt/homebrew/bin/tesseract',  # Mac avec Homebrew ARM
        ]
    
    for path in default_paths:
        if os.path.exists(path):
            return path
    
    return None

# OCR avec Tesseract
try:
    import pytesseract
    from PIL import Image
    
    tesseract_path = find_tesseract()
    if tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        OCR_AVAILABLE = True
        print(f'[OCR] Tesseract configuré: {tesseract_path}')
    else:
        OCR_AVAILABLE = False
        print('[OCR] Tesseract non trouvé. Définissez TESSERACT_PATH ou installez Tesseract.')
except ImportError:
    OCR_AVAILABLE = False
    print('[OCR] pytesseract non installé - pip install pytesseract Pillow')

# Génération DOCX
try:
    from docx import Document
    from docx.shared import Pt, Inches
    DOCX_AVAILABLE = True
    print('[DOCX] python-docx disponible')
except ImportError:
    DOCX_AVAILABLE = False
    print('[DOCX] python-docx non installé - pip install python-docx')

# Génération PDF (modèle caution location)
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
    print('[PDF] reportlab disponible')
except ImportError:
    PDF_AVAILABLE = False
    print('[PDF] reportlab non installé - pip install reportlab')

# PyPDF2 pour fusionner avec le modèle
try:
    from PyPDF2 import PdfReader, PdfWriter
    PYPDF2_AVAILABLE = True
    print('[PDF] PyPDF2 disponible')
except ImportError:
    PYPDF2_AVAILABLE = False
    print('[PDF] PyPDF2 non installé - pip install PyPDF2')

# FPDF pour génération simple de PDF
try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
    print('[PDF] fpdf disponible')
except ImportError:
    FPDF_AVAILABLE = False
    print('[PDF] fpdf non installé - pip install fpdf')

app = Flask(__name__, static_folder='.')

# Configuration pour gérer les URLs longues (éviter erreur 414)
# Augmenter la limite de taille de requête pour les données POST
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB

# ==================== LOGS ERREURS (terminal) ====================
def _log(level, msg, exc=None):
    """Afficher un log avec horodatage dans le terminal (sécurisé Windows)."""
    try:
        ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        line = f"[{ts}] [{level}] {msg}"
        safe_print(line)
    except:
        pass  # Ignorer silencieusement les erreurs de log

# Gestionnaire d'erreur pour les URLs trop longues (414)
@app.errorhandler(414)
def request_uri_too_large(error):
    """Gérer les erreurs 414 (Request-URI Too Long)"""
    _log('ERREUR', f'URL trop longue (414): {request.url[:200]}...')
    return jsonify({
        'success': False,
        'error': 'URL trop longue. Veuillez utiliser POST pour envoyer des données volumineuses.'
    }), 414

@app.errorhandler(500)
def internal_error(error):
    """Toute erreur 500 : log + traceback dans le terminal."""
    _log('ERREUR', f'Erreur 500: {sanitize_error(error)}', error)
    return jsonify({'success': False, 'error': sanitize_error(error)}), 500

@app.errorhandler(Exception)
def handle_exception(error):
    """Exception non gérée : log + traceback dans le terminal."""
    _log('ERREUR', f'Exception: {sanitize_error(error)}', error)
    return jsonify({'success': False, 'error': sanitize_error(error)}), 500

@app.before_request
def log_request():
    """Log chaque requête API dans le terminal pour le debug."""
    if request.path.startswith('/api'):
        _log('API', f"{request.method} {request.path}")

# Configuration CORS sécurisée
CORS(app, 
     resources={r"/api/*": {
         "origins": CORS_ORIGINS_LIST,
         "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
         "allow_headers": ["Content-Type", "Authorization"],
         "expose_headers": ["Content-Type"],
         "max_age": 3600
     }},
     supports_credentials=True if APP_MODE == 'production' else False)

# Système de broadcast pour Server-Sent Events
clients = []
clients_lock = threading.Lock()

def broadcast_event(event_type, data):
    """Diffuser un événement à tous les clients connectés"""
    message = f"data: {json.dumps({'type': event_type, 'data': data})}\n\n"
    
    
    with clients_lock:
        client_count = len(clients)
        
        disconnected_clients = []
        for client_queue in clients:
            try:
                client_queue.put_nowait(message)
            except queue.Full:
                # Queue pleine, ignorer
                pass
            except:
                # Client déconnecté, le marquer pour suppression
                disconnected_clients.append(client_queue)
        
        # Retirer les clients déconnectés
        for client_queue in disconnected_clients:
            if client_queue in clients:
                clients.remove(client_queue)
        
        success_count = client_count - len(disconnected_clients)
        if success_count > 0:
            safe_print(f'[SSE] Event {event_type} -> {success_count} client(s)')

# Configuration base de données (utiliser chemin absolu par défaut)
DB_PATH = os.environ.get('DB_PATH', os.path.join(SCRIPT_DIR, 'data', 'inventory.db'))
print(f'[CONFIG] DB Path: {DB_PATH}')

def sanitize_string(value, max_length=500):
    """Nettoyer et limiter la longueur d'une chaîne, en gérant les caractères Unicode"""
    if value is None:
        return None
    if not value:
        return value
    try:
        # Convertir en string si ce n'est pas déjà le cas
        if not isinstance(value, str):
            value = str(value)
        
        # Supprimer les caractères de contrôle dangereux
        value = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', value)
        
        # Remplacer les caractères Unicode problématiques par des équivalents ASCII
        replacements = {
            '→': '->',
            '←': '<-',
            '↑': '^',
            '↓': 'v',
            '…': '...',
            '–': '-',
            '—': '-',
            '"': '"',
            '"': '"',
            ''': "'",
            ''': "'",
        }
        
        for unicode_char, ascii_replacement in replacements.items():
            value = value.replace(unicode_char, ascii_replacement)
        
        # S'assurer que la chaîne peut être encodée en UTF-8
        try:
            value = value.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
        except:
            # Si l'encodage UTF-8 échoue, utiliser ASCII comme fallback
            value = value.encode('ascii', errors='replace').decode('ascii', errors='replace')
        
        # Limiter la longueur
        return value[:max_length].strip()
    except Exception as e:
        # En cas d'erreur, retourner un message par défaut
        try:
            error_msg = str(e)
            error_msg = error_msg.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
            print(f'[ERREUR] Erreur lors de la sanitization: {error_msg}')
        except:
            print('[ERREUR] Erreur lors de la sanitization (encodage impossible)')
        return 'Message'

def sanitize_error(error):
    """Nettoyer un message d'erreur pour éviter les problèmes d'encodage Unicode"""
    try:
        if isinstance(error, Exception):
            error_str = str(error)
        else:
            error_str = str(error)
        # S'assurer que le message peut être encodé en UTF-8
        try:
            error_str = error_str.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
        except:
            pass
        return sanitize_string(error_str)
    except:
        return 'Une erreur est survenue'

def sanitize_notification_message(message):
    """Nettoyer un message de notification pour éviter les problèmes d'encodage"""
    return sanitize_string(message)

def clean_existing_notifications():
    """Nettoyer toutes les notifications existantes dans la base de données"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        # Récupérer toutes les notifications
        cursor.execute('SELECT id, message FROM notifications')
        rows = cursor.fetchall()
        
        updated_count = 0
        for row in rows:
            original_message = row['message']
            if original_message:
                clean_message = sanitize_notification_message(original_message)
                # Si le message a changé, mettre à jour
                if clean_message != original_message:
                    cursor.execute('''
                        UPDATE notifications 
                        SET message = ? 
                        WHERE id = ?
                    ''', (clean_message, row['id']))
                    updated_count += 1
        
        if updated_count > 0:
            conn.commit()
            print(f'[DB] {updated_count} notification(s) nettoyée(s) pour éviter les problèmes d\'encodage')
        
        conn.close()
    except Exception as e:
        print(f'[DB] Erreur lors du nettoyage des notifications: {str(e)}')

def create_notification(message, type, item_serial_number, conn, cursor):
    """Créer une notification dans la base de données (avec item_hex_id pour navigation)"""
    try:
        # Récupérer l'ID hexadécimal de l'item pour le lien depuis la notification
        item_hex_id = None
        if item_serial_number:
            cursor.execute('SELECT hex_id FROM items WHERE serial_number = ?', (item_serial_number,))
            row = cursor.fetchone()
            if row and row['hex_id']:
                item_hex_id = row['hex_id']
            else:
                # Backfill hex_id pour cet item
                new_hex = generate_item_hex_id(cursor)
                cursor.execute('UPDATE items SET hex_id = ? WHERE serial_number = ?', (new_hex, item_serial_number))
                item_hex_id = new_hex
        # Nettoyer le message pour éviter les problèmes d'encodage
        clean_message = sanitize_notification_message(message)
        now = datetime.now().isoformat()
        cursor.execute('''
            INSERT INTO notifications (message, type, item_serial_number, item_hex_id, created_at)
            VALUES (?, ?, ?, ?, ?)
        ''', (
            clean_message,
            type,
            item_serial_number,
            item_hex_id,
            now
        ))
        notification_id = cursor.lastrowid
        # Ne pas afficher le message dans la console car il peut contenir des caractères Unicode
        # qui causent des erreurs d'encodage sur Windows
        safe_print(f'[API] Notification créée (ID: {notification_id})')
        
        # Limiter à 100 notifications (supprimer les plus anciennes)
        cursor.execute('''
            DELETE FROM notifications 
            WHERE id NOT IN (
                SELECT id FROM notifications 
                ORDER BY created_at DESC 
                LIMIT 100
            )
        ''')
        deleted_count = cursor.rowcount
        if deleted_count > 0:
            safe_print(f'[API] {deleted_count} anciennes notifications supprimées')
    except Exception as e:
        safe_print(f'[API] Erreur lors de la création de la notification: {str(e)}')

def get_db():
    """Créer une connexion à la base de données"""
    os.makedirs(os.path.join(SCRIPT_DIR, 'data'), exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def generate_next_item_id(cursor):
    """Générer le prochain ID alphanumérique (aaa, aab, ..., aaz, aa0, aa1, ..., aa9, aba, etc.)"""
    # Récupérer le dernier item_id utilisé
    cursor.execute('SELECT item_id FROM items WHERE item_id IS NOT NULL ORDER BY item_id DESC LIMIT 1')
    last_id_row = cursor.fetchone()
    
    if not last_id_row or not last_id_row['item_id']:
        # Premier ID : aaa
        return 'aaa'
    
    last_id = last_id_row['item_id'].lower()
    
    # S'assurer que l'ID fait exactement 3 caractères
    if len(last_id) != 3:
        return 'aaa'
    
    # Caractères valides : a-z (26) puis 0-9 (10) = 36 caractères
    chars = 'abcdefghijklmnopqrstuvwxyz0123456789'
    
    # Convertir l'ID en liste de caractères
    id_chars = list(last_id)
    
    # Incrémenter de droite à gauche (position 2 -> 0)
    for i in range(2, -1, -1):
        char_index = chars.find(id_chars[i])
        if char_index == -1:
            # Caractère invalide, réinitialiser à 'a'
            id_chars[i] = 'a'
            continue
        
        if char_index < len(chars) - 1:
            # Incrémenter ce caractère
            id_chars[i] = chars[char_index + 1]
            # Réinitialiser tous les caractères à droite à 'a'
            for j in range(i + 1, 3):
                id_chars[j] = 'a'
            return ''.join(id_chars)
        else:
            # Ce caractère est '9' (dernier), le réinitialiser à 'a' et continuer avec le suivant
            id_chars[i] = 'a'
    
    # Si tous les caractères étaient '9', on recommence à aaa
    # (ne devrait jamais arriver avec seulement 3 caractères)
    return 'aaa'

def generate_item_hex_id(cursor):
    """Générer un ID alphanumérique unique pour un item (format: A00-Z99)
    A00, A01, ... A99, B00, B01, ... Z99 = 2600 IDs possibles
    """
    # Chercher le dernier hex_id utilisé (nouveau format A00-Z99)
    cursor.execute('''
        SELECT hex_id FROM items 
        WHERE hex_id IS NOT NULL AND length(hex_id) = 3 
        AND hex_id GLOB '[A-Z][0-9][0-9]'
        ORDER BY hex_id DESC LIMIT 1
    ''')
    row = cursor.fetchone()
    
    if row and row['hex_id']:
        try:
            last_id = row['hex_id']
            letter = last_id[0]
            number = int(last_id[1:3])
            
            # Incrémenter
            number += 1
            if number > 99:
                number = 0
                # Passer à la lettre suivante
                if letter == 'Z':
                    letter = 'A'  # Retour au début (ne devrait jamais arriver avec 2600 IDs)
                else:
                    letter = chr(ord(letter) + 1)
            
            return f"{letter}{number:02d}"
        except (ValueError, IndexError):
            pass
    
    # Si pas de hex_id existant ou erreur, commencer à A00
    return 'A00'

def migrate_hex_ids():
    """Migrer tous les hex_id vers le nouveau format alphanumérique (A00-Z99)"""
    conn = get_db()
    cursor = conn.cursor()
    
    try:
        # Récupérer tous les items qui n'ont pas le nouveau format (A00-Z99)
        cursor.execute('''
            SELECT id FROM items 
            WHERE hex_id IS NULL 
               OR hex_id NOT GLOB '[A-Z][0-9][0-9]'
            ORDER BY id ASC
        ''')
        items_to_update = cursor.fetchall()
        
        if items_to_update:
            print(f'[DB] Migration de {len(items_to_update)} hex_id vers format A00-Z99...')
            
            # Trouver le dernier ID utilisé dans le nouveau format
            cursor.execute('''
                SELECT hex_id FROM items 
                WHERE hex_id GLOB '[A-Z][0-9][0-9]'
                ORDER BY hex_id DESC LIMIT 1
            ''')
            last_row = cursor.fetchone()
            
            if last_row and last_row['hex_id']:
                letter = last_row['hex_id'][0]
                number = int(last_row['hex_id'][1:3])
            else:
                letter = 'A'
                number = -1  # Commencera à 0 après incrémentation
            
            for item_row in items_to_update:
                item_id = item_row['id']
                # Incrémenter
                number += 1
                if number > 99:
                    number = 0
                    letter = chr(ord(letter) + 1) if letter != 'Z' else 'A'
                
                new_id = f"{letter}{number:02d}"
                cursor.execute('UPDATE items SET hex_id = ? WHERE id = ?', (new_id, item_id))
            
            conn.commit()
            print(f'[DB] Migration hex_id terminée: {len(items_to_update)} items mis à jour')
    except Exception as e:
        print(f'[DB] Erreur migration hex_id: {str(e)}')
    finally:
        conn.close()

def init_db():
    """Initialiser la base de données"""
    conn = get_db()
    cursor = conn.cursor()
    
    # Table des items
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            item_id TEXT,
            name TEXT NOT NULL,
            serial_number TEXT NOT NULL UNIQUE,
            quantity INTEGER DEFAULT 1,
            category TEXT,
            category_details TEXT,
            image TEXT,
            scanned_code TEXT,
            created_at TEXT NOT NULL,
            last_updated TEXT NOT NULL
        )
    ''')
    
    # Vérifier et ajouter les colonnes manquantes pour les bases de données existantes
    cursor.execute("PRAGMA table_info(items)")
    columns = [column[1] for column in cursor.fetchall()]
    print(f'[DB] Colonnes existantes dans items: {columns}')
    
    # Ajouter les nouvelles colonnes si elles n'existent pas
    # SQLite n'accepte pas UNIQUE/DEFAULT complexes dans ALTER ADD COLUMN, donc on utilise seulement le type
    new_columns = {
        'item_id': 'TEXT',
        'hex_id': 'TEXT',  # ID hexadécimal unique pour navigation (notifications)
        'status': 'TEXT',
        'item_type': 'TEXT',
        'brand': 'TEXT',
        'model': 'TEXT',
        'rental_end_date': 'TEXT',
        'current_rental_id': 'INTEGER',
        'custom_data': 'TEXT',  # JSON pour stocker les champs personnalisés
        'parent_id': 'INTEGER',  # ID du parent pour créer des groupes d'items
        'display_order': 'INTEGER'  # Ordre d'affichage dans le groupe
    }
    
    for col_name, col_type in new_columns.items():
        if col_name not in columns:
            try:
                cursor.execute(f'ALTER TABLE items ADD COLUMN {col_name} {col_type}')
                print(f'[DB] Colonne {col_name} ajoutée')
            except sqlite3.OperationalError as e:
                print(f'[DB] Erreur ajout colonne {col_name}: {e}')
    
    conn.commit()
    
    # Table des champs personnalisés
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS custom_fields (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            field_key TEXT NOT NULL UNIQUE,
            field_type TEXT NOT NULL DEFAULT 'text',
            options TEXT,
            required INTEGER DEFAULT 0,
            display_order INTEGER DEFAULT 0,
            created_at TEXT NOT NULL
        )
    ''')
    conn.commit()
    
    # Table des catégories personnalisées
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS custom_categories (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            created_at TEXT NOT NULL
        )
    ''')
    
    # Table des catégories supprimées
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS deleted_categories (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            deleted_at TEXT NOT NULL
        )
    ''')
    
    # Initialiser les nouvelles catégories d'équipement
    new_equipment_categories = ['ordinateur', 'casque_vr', 'camera', 'eclairage', 'accessoire']
    for cat_name in new_equipment_categories:
        try:
            cursor.execute(
                'INSERT OR IGNORE INTO custom_categories (name, created_at) VALUES (?, ?)',
                (cat_name, datetime.now().isoformat())
            )
        except sqlite3.IntegrityError:
            pass  # La catégorie existe déjà
    
    conn.commit()
    
    # Table d'historique des modifications d'items
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS item_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            item_serial_number TEXT NOT NULL,
            field_name TEXT NOT NULL,
            old_value TEXT,
            new_value TEXT,
            changed_at TEXT NOT NULL,
            FOREIGN KEY (item_serial_number) REFERENCES items(serial_number)
        )
    ''')
    
    # Table des notifications partagées
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS notifications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            message TEXT NOT NULL,
            type TEXT NOT NULL,
            item_serial_number TEXT,
            item_hex_id TEXT,
            created_at TEXT NOT NULL
        )
    ''')
    # Ajouter item_hex_id si la table existait déjà
    cursor.execute("PRAGMA table_info(notifications)")
    notif_columns = [c[1] for c in cursor.fetchall()]
    if 'item_hex_id' not in notif_columns:
        try:
            cursor.execute('ALTER TABLE notifications ADD COLUMN item_hex_id TEXT')
        except sqlite3.OperationalError:
            pass
    
    conn.commit()
    
    # Nettoyer les notifications existantes pour éviter les problèmes d'encodage
    clean_existing_notifications()
    
    # Table des locations
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS rentals (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            renter_name TEXT NOT NULL,
            renter_email TEXT NOT NULL,
            renter_phone TEXT NOT NULL,
            renter_address TEXT,
            rental_price REAL NOT NULL,
            rental_deposit REAL NOT NULL,
            rental_duration INTEGER NOT NULL,
            start_date TEXT NOT NULL,
            end_date TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'en_cours',
            items_data TEXT NOT NULL,
            attachments TEXT,
            notes TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
    ''')
    
    # Ajouter colonnes manquantes à rentals
    cursor.execute("PRAGMA table_info(rentals)")
    rental_columns = [column[1] for column in cursor.fetchall()]
    
    rental_new_columns = {
        'attachments': 'TEXT',
        'notes': 'TEXT'
    }
    
    for col_name, col_type in rental_new_columns.items():
        if col_name not in rental_columns:
            try:
                cursor.execute(f'ALTER TABLE rentals ADD COLUMN {col_name} {col_type}')
                print(f'[DB] Colonne rentals.{col_name} ajoutée')
            except sqlite3.OperationalError as e:
                print(f'[DB] Erreur ajout colonne rentals.{col_name}: {e}')
    
    # Table des statuts personnalisés pour les locations
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS rental_statuses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            color TEXT DEFAULT '#666',
            created_at TEXT NOT NULL
        )
    ''')
    
    # Insérer les statuts par défaut
    default_statuses = [
        ('en_cours', '#007bff'),
        ('contrat_envoye', '#ffc107'),
        ('fini', '#28a745')
    ]
    for status_name, color in default_statuses:
        cursor.execute('''
            INSERT OR IGNORE INTO rental_statuses (name, color, created_at)
            VALUES (?, ?, ?)
        ''', (status_name, color, datetime.now().isoformat()))
    
    # Index pour améliorer les performances
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_item_history_serial ON item_history(item_serial_number)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_notifications_created_at ON notifications(created_at)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_rentals_start_date ON rentals(start_date)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_rentals_end_date ON rentals(end_date)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_rentals_status ON rentals(status)')
    
    conn.commit()
    conn.close()
    print(f"[OK] Base de donnees initialisee: {DB_PATH}")

# ==================== CONFIGURATION FRONTEND STATIQUE ====================

# Chemin vers le build du frontend Next.js (tout à la racine du projet)
FRONTEND_BUILD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'out')

def check_frontend_build():
    """Vérifier si le frontend est buildé"""
    if os.path.exists(FRONTEND_BUILD_DIR) and os.path.isdir(FRONTEND_BUILD_DIR):
        # Vérifier qu'il y a au moins un fichier index.html
        index_path = os.path.join(FRONTEND_BUILD_DIR, 'index.html')
        return os.path.exists(index_path)
    return False

FRONTEND_AVAILABLE = check_frontend_build()
if FRONTEND_AVAILABLE:
    print(f'[FRONTEND] Build trouvé dans: {FRONTEND_BUILD_DIR}')
else:
    print(f'[FRONTEND] Build non trouvé. Exécutez: npm run build')

# ==================== ROUTES FRONTEND (fichiers statiques) ====================

@app.route('/')
def serve_index():
    """Servir la page d'accueil du frontend"""
    if FRONTEND_AVAILABLE:
        return send_from_directory(FRONTEND_BUILD_DIR, 'index.html')
    else:
        return '''
        <html>
        <head><title>Code Bar CRM - Build requis</title></head>
        <body style="font-family: sans-serif; text-align: center; padding: 50px;">
            <h1>🔧 Frontend non buildé</h1>
            <p>Pour utiliser l'application, vous devez d'abord compiler le frontend :</p>
            <pre style="background: #f4f4f4; padding: 20px; display: inline-block; text-align: left;">
npm install
npm run build
            </pre>
            <p>Puis relancez le serveur :</p>
            <pre style="background: #f4f4f4; padding: 20px; display: inline-block;">python server.py</pre>
            <hr>
            <p>L'API est disponible sur <code>/api/*</code></p>
        </body>
        </html>
        ''', 200

@app.route('/favicon.ico')
def serve_favicon():
    """Servir le favicon"""
    if FRONTEND_AVAILABLE:
        favicon_path = os.path.join(FRONTEND_BUILD_DIR, 'favicon.ico')
        if os.path.exists(favicon_path):
            return send_from_directory(FRONTEND_BUILD_DIR, 'favicon.ico')
    return '', 204

@app.route('/logo-globalvision.png')
def serve_logo():
    """Servir le logo"""
    if os.path.exists('logo-globalvision.png'):
        return send_from_directory('.', 'logo-globalvision.png')
    return '', 404

# Route pour servir les fichiers statiques Next.js (_next, images, etc.)
@app.route('/_next/<path:filename>')
def serve_next_static(filename):
    """Servir les fichiers statiques Next.js (_next)"""
    if FRONTEND_AVAILABLE:
        return send_from_directory(os.path.join(FRONTEND_BUILD_DIR, '_next'), filename)
    return '', 404

@app.route('/fonts/<path:filename>')
def serve_fonts(filename):
    """Servir les polices"""
    if FRONTEND_AVAILABLE:
        fonts_dir = os.path.join(FRONTEND_BUILD_DIR, 'fonts')
        if os.path.exists(fonts_dir):
            return send_from_directory(fonts_dir, filename)
    return '', 404

@app.route('/img/<path:filename>')
def serve_images(filename):
    """Servir les images"""
    if FRONTEND_AVAILABLE:
        img_dir = os.path.join(FRONTEND_BUILD_DIR, 'img')
        if os.path.exists(img_dir):
            return send_from_directory(img_dir, filename)
    return '', 404

# ==================== API IMAGES ====================

@app.route('/api/images/<filename>')
def serve_image(filename):
    """Servir une image sauvegardée"""
    try:
        # Sécurité: empêcher la traversée de répertoire
        safe_filename = os.path.basename(filename)
        return send_from_directory(IMAGES_DIR, safe_filename)
    except Exception as e:
        safe_print(f'[IMG] Erreur lecture image {filename}: {str(e)}')
        return jsonify({'error': 'Image non trouvée'}), 404

@app.route('/api/upload-image', methods=['POST'])
def upload_image():
    """Upload une image directement (sans Base64) - Simple et fiable"""
    try:
        safe_print(f'[IMG] Upload request - files: {list(request.files.keys())}')
        
        if 'file' not in request.files:
            safe_print('[IMG] Erreur: pas de fichier dans la requete')
            return jsonify({'success': False, 'error': 'Aucun fichier'}), 400
        
        file = request.files['file']
        if not file or not file.filename:
            safe_print('[IMG] Erreur: fichier vide')
            return jsonify({'success': False, 'error': 'Fichier vide'}), 400
        
        safe_print(f'[IMG] Fichier recu: {file.filename}')
        
        # Récupérer le serial_number optionnel
        serial_number = request.form.get('serialNumber', 'img')
        
        # Sauvegarder le fichier
        path = save_uploaded_file(file, serial_number)
        if path:
            safe_print(f'[IMG] Upload reussi: {path}')
            return jsonify({'success': True, 'path': path}), 200
        else:
            safe_print('[IMG] Erreur: save_uploaded_file a retourne None')
            return jsonify({'success': False, 'error': 'Erreur sauvegarde fichier'}), 500
            
    except Exception as e:
        safe_print(f'[IMG] Exception upload: {e}')
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

# ==================== API ITEMS ====================

@app.route('/api/items', methods=['GET'])
def get_items():
    """Récupérer tous les items"""
    conn = None
    try:
        print('[API] GET /api/items - Récupération des items...')
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM items ORDER BY last_updated DESC')
        rows = cursor.fetchall()
        
        items = []
        for row in rows:
            # Convertir le Row en dict pour accès sécurisé
            row_dict = dict(row)
            # Backfill hex_id si manquant (ID hexadécimal unique pour notifications/navigation)
            hex_id = row_dict.get('hex_id')
            if not hex_id:
                hex_id = generate_item_hex_id(cursor)
                cursor.execute('UPDATE items SET hex_id = ? WHERE id = ?', (hex_id, row_dict['id']))
                conn.commit()
            # Parser custom_data JSON
            custom_data = {}
            if row_dict.get('custom_data'):
                try:
                    custom_data = json.loads(row_dict['custom_data'])
                except:
                    pass
            
            items.append({
                'id': row_dict.get('id'),
                'itemId': row_dict.get('item_id'),
                'hexId': hex_id,
                'name': row_dict.get('name'),
                'serialNumber': row_dict.get('serial_number'),
                'quantity': row_dict.get('quantity', 1),
                'category': row_dict.get('category'),
                'categoryDetails': row_dict.get('category_details'),
                'image': row_dict.get('image'),
                'scannedCode': row_dict.get('scanned_code'),
                'status': row_dict.get('status', 'en_stock'),
                'itemType': row_dict.get('item_type'),
                'brand': row_dict.get('brand'),
                'model': row_dict.get('model'),
                'rentalEndDate': row_dict.get('rental_end_date'),
                'currentRentalId': row_dict.get('current_rental_id'),
                'parentId': row_dict.get('parent_id'),
                'displayOrder': row_dict.get('display_order', 0),
                'customData': custom_data,
                'createdAt': row_dict.get('created_at'),
                'lastUpdated': row_dict.get('last_updated')
            })
        
        if conn:
            conn.close()
        print(f'[API] GET /api/items - {len(items)} items retournés')
        return jsonify({'success': True, 'items': items}), 200
    except Exception as e:
        if conn:
            try:
                conn.close()
            except Exception:
                pass
        print(f'[API] ERREUR GET /api/items: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/items/search', methods=['GET'])
def search_item():
    """Rechercher un item par numéro de série ou code-barres"""
    try:
        query = request.args.get('q', '').strip()
        if not query:
            return jsonify({'success': False, 'error': 'Paramètre de recherche manquant'}), 400
        
        print(f'[API] GET /api/items/search - Recherche: {query}')
        conn = get_db()
        cursor = conn.cursor()
        
        # Rechercher par numéro de série, code scanné, item_id ou hex_id
        cursor.execute('''
            SELECT * FROM items 
            WHERE serial_number = ? OR scanned_code = ? OR item_id = ? OR hex_id = ?
            LIMIT 1
        ''', (query, query, query, query))
        
        row = cursor.fetchone()
        conn.close()
        
        if row:
            row_dict = dict(row)
            hex_id = row_dict.get('hex_id')
            if not hex_id:
                conn = get_db()
                cur = conn.cursor()
                hex_id = generate_item_hex_id(cur)
                cur.execute('UPDATE items SET hex_id = ? WHERE id = ?', (hex_id, row_dict['id']))
                conn.commit()
                conn.close()
            item = {
                'id': row_dict.get('id'),
                'itemId': row_dict.get('item_id'),
                'hexId': hex_id,
                'name': row_dict.get('name'),
                'serialNumber': row_dict.get('serial_number'),
                'quantity': row_dict.get('quantity', 1),
                'category': row_dict.get('category'),
                'categoryDetails': row_dict.get('category_details'),
                'image': row_dict.get('image'),
                'scannedCode': row_dict.get('scanned_code'),
                'status': row_dict.get('status', 'en_stock'),
                'itemType': row_dict.get('item_type'),
                'brand': row_dict.get('brand'),
                'model': row_dict.get('model'),
                'rentalEndDate': row_dict.get('rental_end_date'),
                'currentRentalId': row_dict.get('current_rental_id'),
                'parentId': row_dict.get('parent_id'),
                'displayOrder': row_dict.get('display_order', 0),
                'createdAt': row_dict.get('created_at'),
                'lastUpdated': row_dict.get('last_updated')
            }
            print(f'[API] GET /api/items/search - Item trouvé: {item["name"]}')
            return jsonify({'success': True, 'found': True, 'item': item}), 200
        else:
            print(f'[API] GET /api/items/search - Aucun item trouvé')
            return jsonify({'success': True, 'found': False, 'item': None}), 200
            
    except Exception as e:
        print(f'[API] ERREUR GET /api/items/search: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/items', methods=['POST'])
def create_item():
    """Créer ou mettre à jour un item"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Données JSON invalides'}), 400

        # Log sans l'image pour éviter [Errno 22] sur Windows (image Base64 trop grande)
        try:
            log_data = {k: (v[:100] + '...' if k == 'image' and isinstance(v, str) and len(v) > 100 else v) for k, v in data.items()}
            safe_print(f'[API] POST /api/items - Données reçues: {log_data}')
        except Exception as log_err:
            safe_print(f'[API] POST /api/items - (log error: {log_err})')
        
        # Validation des champs requis
        missing_fields = validate_required_fields(data, ['name', 'serialNumber'])
        if missing_fields:
            safe_print(f'[API] ERREUR: Champs manquants: {missing_fields}')
            return jsonify({'success': False, 'error': f'Champs obligatoires manquants: {", ".join(missing_fields)}'}), 400
        
        # Sanitization des données
        data['name'] = sanitize_string(data.get('name'), 200)
        data['serialNumber'] = sanitize_string(data.get('serialNumber'), 100)
        data['brand'] = sanitize_string(data.get('brand'), 100)
        data['model'] = sanitize_string(data.get('model'), 100)
        data['category'] = sanitize_string(data.get('category'), 50)
        data['categoryDetails'] = sanitize_string(data.get('categoryDetails'), 1000)
        
        # Validation de la quantité
        quantity = data.get('quantity', 1)
        if not validate_positive_number(quantity, allow_zero=False):
            return jsonify({'success': False, 'error': 'La quantité doit être un nombre positif'}), 400
        data['quantity'] = int(quantity)
        
        if not data.get('name') or not data.get('serialNumber'):
            safe_print('[API] ERREUR: Nom ou numéro de série manquant après sanitization')
            return jsonify({'success': False, 'error': 'Le nom et le numéro de série sont obligatoires'}), 400

        # Traiter les images: sauvegarder sur disque au lieu de stocker en Base64
        if data.get('image'):
            try:
                safe_print('[API] Traitement des images...')
                data['image'] = process_images_for_storage(data['image'], data['serialNumber'])
                safe_print(f'[API] Images traitées: {str(data.get("image", "None"))[:100]}')
            except Exception as img_err:
                safe_print(f'[API] Erreur traitement images (ignorée): {str(img_err)}')
                data['image'] = None  # Continuer sans images

        now = datetime.now().isoformat()
        conn = get_db()
        cursor = conn.cursor()
        
        # Vérifier si l'item existe déjà
        cursor.execute('SELECT * FROM items WHERE serial_number = ?', (data['serialNumber'],))
        existing = cursor.fetchone()
        
        if existing:
            safe_print(f'[API] Item existant trouvé (ID: {existing["id"]}), mise à jour...')
            # Mettre à jour l'item existant (ajouter la quantité)
            quantity_to_add = data.get('quantity', 1)
            old_quantity = existing['quantity']
            new_quantity = old_quantity + quantity_to_add
            
            # Enregistrer l'historique de la modification de quantité
            if quantity_to_add > 0:
                cursor.execute('''
                    INSERT INTO item_history (item_serial_number, field_name, old_value, new_value, changed_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    data['serialNumber'],
                    'quantity',
                    str(old_quantity),
                    str(new_quantity),
                    now
                ))
            
            # Préparer custom_data
            custom_data = data.get('customData')
            custom_data_json = json.dumps(custom_data) if custom_data else None
            
            cursor.execute('''
                UPDATE items 
                SET name = ?, quantity = ?, category = ?, category_details = ?, 
                    image = ?, scanned_code = ?, item_type = ?, brand = ?, model = ?, custom_data = ?, last_updated = ?
                WHERE serial_number = ?
            ''', (
                data['name'],
                new_quantity,
                data.get('category'),
                data.get('categoryDetails'),
                data.get('image'),
                data.get('scannedCode') or data['serialNumber'],
                data.get('itemType'),
                data.get('brand'),
                data.get('model'),
                custom_data_json,
                now,
                data['serialNumber']
            ))
            item_id = existing['id']
            safe_print(f'[API] Item mis à jour (ID: {item_id}, quantité: {new_quantity})')
        else:
            safe_print('[API] Nouvel item, création...')
            # Générer un nouvel item_id et un ID hexadécimal unique
            item_id_code = generate_next_item_id(cursor)
            hex_id = generate_item_hex_id(cursor)
            safe_print(f'[API] Nouvel item_id généré: {item_id_code}, hex_id: {hex_id}')
            
            # Préparer custom_data pour nouvel item
            custom_data = data.get('customData')
            custom_data_json = json.dumps(custom_data) if custom_data else None
            
            # Créer un nouvel item
            cursor.execute('''
                INSERT INTO items (item_id, hex_id, name, serial_number, quantity, category, category_details, 
                                 image, scanned_code, item_type, brand, model, status, custom_data, created_at, last_updated)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                item_id_code,
                hex_id,
                data['name'],
                data['serialNumber'],
                data.get('quantity', 1),
                data.get('category'),
                data.get('categoryDetails'),
                data.get('image'),
                data.get('scannedCode') or data['serialNumber'],
                data.get('itemType'),
                data.get('brand'),
                data.get('model'),
                'en_stock',
                custom_data_json,
                now,
                now
            ))
            item_id = cursor.lastrowid
            
            # Enregistrer la création dans l'historique
            cursor.execute('''
                INSERT INTO item_history (item_serial_number, field_name, old_value, new_value, changed_at)
                VALUES (?, ?, ?, ?, ?)
            ''', (
                data['serialNumber'],
                'created',
                None,
                'Item créé',
                now
            ))
            
            safe_print(f'[API] Nouvel item créé (ID: {item_id})')

        # Créer une notification avec heure (format simplifié pour éviter Errno 22 sur Windows)
        try:
            now_dt = datetime.now()
            time_str = f"{now_dt.hour:02d}:{now_dt.minute:02d}:{now_dt.second:02d}"
            date_str = f"{now_dt.day:02d}/{now_dt.month:02d}/{now_dt.year}"
        except Exception as dt_err:
            safe_print(f'[API] Erreur date/heure: {dt_err}')
            time_str = "00:00:00"
            date_str = "01/01/2025"
        
        try:
            if existing:
                create_notification(
                    f'Modification de quantite - Item "{data["name"]}" ({data["serialNumber"]}) : {old_quantity} -> {new_quantity} | {date_str} {time_str}',
                    'success',
                    data['serialNumber'],
                    conn,
                    cursor
                )
            else:
                create_notification(
                    f'Nouvel item cree - "{data["name"]}" ({data["serialNumber"]}) | {date_str} {time_str}',
                    'success',
                    data['serialNumber'],
                    conn,
                    cursor
                )
        except Exception as notif_err:
            safe_print(f'[API] Erreur notification (non bloquante): {notif_err}')
        
        try:
            conn.commit()
        except Exception as commit_err:
            safe_print(f'[API] Erreur commit DB: {commit_err}')
            raise commit_err
        finally:
            conn.close()
        
        # Diffuser l'événement à tous les clients
        broadcast_event('items_changed', {'action': 'created' if not existing else 'updated', 'id': item_id})
        broadcast_event('notifications_changed', {})

        safe_print(f'[API] POST /api/items - Succès (ID: {item_id})')
        return jsonify({'success': True, 'id': item_id}), 201
    except Exception as e:
        safe_print(f'[API] ERREUR POST /api/items: {str(e)}')
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/items/<serial_number>', methods=['PUT'])
def update_item(serial_number):
    """Mettre à jour un item"""
    try:
        data = request.get_json()

        # Traiter les images si présentes
        if data.get('image'):
            try:
                safe_print('[API] PUT - Traitement des images...')
                data['image'] = process_images_for_storage(data['image'], serial_number)
            except Exception as img_err:
                safe_print(f'[API] PUT - Erreur images (ignorée): {str(img_err)}')
                data['image'] = None

        conn = get_db()
        cursor = conn.cursor()

        # Récupérer l'item existant pour comparer les valeurs
        cursor.execute('SELECT * FROM items WHERE serial_number = ?', (serial_number,))
        existing_row = cursor.fetchone()
        if not existing_row:
            conn.close()
            return jsonify({'success': False, 'error': 'Item non trouvé'}), 404
        existing = dict(existing_row)

        now = datetime.now().isoformat()

        # Mapping des champs API vers colonnes DB
        field_mapping = {
            'name': 'name',
            'quantity': 'quantity',
            'category': 'category',
            'categoryDetails': 'category_details',
            'image': 'image',
            'scannedCode': 'scanned_code',
            'serialNumber': 'serial_number',
            'brand': 'brand',
            'model': 'model',
            'itemType': 'item_type',
            'status': 'status'
        }

        # Construire la requête de mise à jour et enregistrer l'historique
        update_fields = []
        update_values = []
        history_entries = []

        for api_field, db_column in field_mapping.items():
            if api_field in data:
                old_value = existing.get(db_column)
                new_value = data[api_field]
                
                # Convertir en string pour la comparaison
                old_val_str = str(old_value) if old_value is not None else None
                new_val_str = str(new_value) if new_value is not None else None
                
                # Enregistrer dans l'historique si la valeur a changé
                if old_val_str != new_val_str:
                    update_fields.append(f'{db_column} = ?')
                    update_values.append(new_value)
                    
                    # Enregistrer dans l'historique
                    history_entries.append({
                        'item_serial_number': serial_number,
                        'field_name': api_field,
                        'old_value': old_val_str,
                        'new_value': new_val_str,
                        'changed_at': now
                    })
        
        # Gérer customData (champs personnalisés)
        if 'customData' in data:
            old_custom_data = {}
            if existing.get('custom_data'):
                try:
                    old_custom_data = json.loads(existing['custom_data'])
                except:
                    pass
            
            new_custom_data = data.get('customData', {})
            
            # Comparer chaque champ personnalisé
            all_custom_keys = set(list(old_custom_data.keys()) + list(new_custom_data.keys()))
            for custom_key in all_custom_keys:
                old_val = old_custom_data.get(custom_key)
                new_val = new_custom_data.get(custom_key)
                
                old_val_str = str(old_val) if old_val is not None else None
                new_val_str = str(new_val) if new_val is not None else None
                
                if old_val_str != new_val_str:
                    history_entries.append({
                        'item_serial_number': serial_number,
                        'field_name': f'custom_{custom_key}',
                        'old_value': old_val_str,
                        'new_value': new_val_str,
                        'changed_at': now
                    })
            
            # Mettre à jour custom_data dans la base
            custom_data_json = json.dumps(new_custom_data) if new_custom_data else None
            update_fields.append('custom_data = ?')
            update_values.append(custom_data_json)
        
        # Si le serialNumber change, mettre à jour la référence dans l'historique
        if 'serialNumber' in data and data['serialNumber'] != serial_number:
            new_serial = data['serialNumber']
            # Mettre à jour les références dans l'historique
            cursor.execute('UPDATE item_history SET item_serial_number = ? WHERE item_serial_number = ?', 
                         (new_serial, serial_number))
        
        if update_fields:
            update_fields.append('last_updated = ?')
            update_values.append(now)
            update_values.append(serial_number)
            
            cursor.execute(
                f'UPDATE items SET {", ".join(update_fields)} WHERE serial_number = ?',
                update_values
            )
            
            # Enregistrer l'historique et créer des notifications
            item_name = existing['name']
            for entry in history_entries:
                cursor.execute('''
                    INSERT INTO item_history (item_serial_number, field_name, old_value, new_value, changed_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    entry['item_serial_number'],
                    entry['field_name'],
                    entry['old_value'],
                    entry['new_value'],
                    entry['changed_at']
                ))
                
                # Créer une notification pour chaque modification
                field_labels = {
                    'name': 'Nom',
                    'quantity': 'Quantité',
                    'category': 'Catégorie',
                    'categoryDetails': 'Détails',
                    'serialNumber': 'Numéro de série',
                    'scannedCode': 'Code scanné',
                    'brand': 'Marque',
                    'model': 'Modèle',
                    'itemType': 'Type',
                    'status': 'Statut',
                    'image': 'Image'
                }
                
                # Gérer les champs personnalisés
                field_name = entry['field_name']
                if field_name.startswith('custom_'):
                    # Récupérer le nom du champ personnalisé depuis la base
                    custom_key = field_name.replace('custom_', '')
                    cursor.execute('SELECT name FROM custom_fields WHERE field_key = ?', (custom_key,))
                    custom_field = cursor.fetchone()
                    field_label = custom_field['name'] if custom_field else custom_key
                else:
                    field_label = field_labels.get(field_name, field_name)
                
                # Formater l'heure complète
                try:
                    changed_time = datetime.fromisoformat(entry['changed_at'].replace('Z', '+00:00'))
                    time_str = changed_time.strftime('%H:%M:%S')
                    date_str = changed_time.strftime('%d/%m/%Y')
                except:
                    changed_time = datetime.now()
                    time_str = changed_time.strftime('%H:%M:%S')
                    date_str = changed_time.strftime('%d/%m/%Y')
                
                old_val_display = entry['old_value'] if entry['old_value'] else 'vide'
                new_val_display = entry['new_value'] if entry['new_value'] else 'vide'
                
                # Formater le message avec tous les détails
                if field_name == 'quantity':
                    notification_msg = f'📊 Modification de {field_label} - Item "{item_name}" ({serial_number}) : {old_val_display} -> {new_val_display} | {date_str} {time_str}'
                elif field_name == 'name':
                    notification_msg = f'✏️ Modification de {field_label} - Item "{old_val_display}" ({serial_number}) -> "{new_val_display}" | {date_str} {time_str}'
                elif field_name == 'category':
                    notification_msg = f'🏷️ Modification de {field_label} - Item "{item_name}" ({serial_number}) : {old_val_display or "aucune"} -> {new_val_display} | {date_str} {time_str}'
                elif field_name == 'status':
                    notification_msg = f'🔄 Modification de {field_label} - Item "{item_name}" ({serial_number}) : {old_val_display} -> {new_val_display} | {date_str} {time_str}'
                elif field_name.startswith('custom_'):
                    notification_msg = f'📝 Modification de {field_label} - Item "{item_name}" ({serial_number}) : {old_val_display} -> {new_val_display} | {date_str} {time_str}'
                else:
                    notification_msg = f'✏️ Modification de {field_label} - Item "{item_name}" ({serial_number}) : {old_val_display} -> {new_val_display} | {date_str} {time_str}'
                
                create_notification(notification_msg, 'success', serial_number, conn, cursor)
        
        conn.commit()
        conn.close()
        
        # Diffuser l'événement à tous les clients
        broadcast_event('items_changed', {'action': 'updated', 'serialNumber': serial_number})
        broadcast_event('notifications_changed', {})
        
        return jsonify({'success': True}), 200
    except Exception as e:
        print(f'[API] ERREUR PUT /api/items/{serial_number}: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/events', methods=['GET'])
def stream_events():
    """Stream Server-Sent Events pour la synchronisation en temps réel"""
    def event_stream():
        # Créer une queue pour ce client (taille limitée pour éviter l'accumulation)
        client_queue = queue.Queue(maxsize=10)
        
        with clients_lock:
            clients.append(client_queue)
            client_count = len(clients)
        
        print(f'[SSE] Client connected. Total: {client_count}')
        
        try:
            # Envoyer un message de connexion
            yield f"data: {json.dumps({'type': 'connected', 'data': {}})}\n\n"
            
            # Garder la connexion ouverte et envoyer les événements
            while True:
                try:
                    # Attendre un événement avec timeout pour vérifier la connexion
                    message = client_queue.get(timeout=30)
                    yield message
                except queue.Empty:
                    # Envoyer un keepalive pour maintenir la connexion
                    yield ": keepalive\n\n"
        except GeneratorExit:
            # Client déconnecté
            pass
        finally:
            # Retirer le client de la liste
            with clients_lock:
                if client_queue in clients:
                    clients.remove(client_queue)
                client_count = len(clients)
            
            print(f'[SSE] Client disconnected. Remaining: {client_count}')
    
    return Response(event_stream(), mimetype='text/event-stream', headers={
        'Cache-Control': 'no-cache',
        'X-Accel-Buffering': 'no',
        'Connection': 'keep-alive'
    })

@app.route('/api/items/<serial_number>/history', methods=['GET'])
def get_item_history(serial_number):
    """Récupérer l'historique des modifications d'un item"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT field_name, old_value, new_value, changed_at
            FROM item_history
            WHERE item_serial_number = ?
            ORDER BY changed_at DESC
            LIMIT 10
        ''', (serial_number,))
        
        rows = cursor.fetchall()
        history = [{
            'fieldName': row['field_name'],
            'oldValue': row['old_value'],
            'newValue': row['new_value'],
            'changedAt': row['changed_at']
        } for row in rows]
        
        conn.close()
        return jsonify({'success': True, 'history': history}), 200
    except Exception as e:
        print(f'[API] ERREUR GET /api/items/{serial_number}/history: {str(e)}')
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/items/delete-all', methods=['POST'])
def delete_all_items():
    """Supprimer tous les items de l'inventaire"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        # Compter les items avant suppression
        cursor.execute('SELECT COUNT(*) as count FROM items')
        count_row = cursor.fetchone()
        item_count = count_row['count'] if count_row else 0
        
        # Supprimer tous les items
        cursor.execute('DELETE FROM items')
        
        # Réinitialiser l'auto-increment
        cursor.execute('DELETE FROM sqlite_sequence WHERE name="items"')
        
        conn.commit()
        
        # Créer une notification
        try:
            time_str = datetime.now().strftime('%H:%M:%S')
            date_str = datetime.now().strftime('%d/%m/%Y')
        except:
            time_str = datetime.now().strftime('%H:%M:%S')
            date_str = datetime.now().strftime('%d/%m/%Y')
        
        create_notification(
            f'🗑️ Inventaire vidé - {item_count} article(s) supprimé(s) | {date_str} {time_str}',
            'warning',
            None,
            conn,
            cursor
        )
        
        conn.commit()
        conn.close()
        
        # Broadcaster la suppression
        broadcast_sse_event('items_changed', {
            'action': 'all_deleted',
            'count': item_count
        })
        broadcast_sse_event('notifications_changed', {})
        
        return jsonify({'success': True, 'count': item_count})
        
    except Exception as e:
        print(f'[API] ERREUR DELETE /api/items/delete-all: {str(e)}')
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/items/<path:serial_number>', methods=['DELETE'])
def delete_item(serial_number):
    """Supprimer un item"""
    try:
        # Décoder l'URL pour gérer les caractères spéciaux comme /
        serial_number = urllib.parse.unquote(serial_number)
        
        conn = get_db()
        cursor = conn.cursor()
        
        # Récupérer le nom de l'item avant suppression
        cursor.execute('SELECT name FROM items WHERE serial_number = ?', (serial_number,))
        item = cursor.fetchone()
        item_name = item['name'] if item else 'Item'
        
        cursor.execute('DELETE FROM items WHERE serial_number = ?', (serial_number,))
        
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({'success': False, 'error': 'Item non trouvé'}), 404
        
        # Créer une notification avec heure
        try:
            time_str = datetime.now().strftime('%H:%M:%S')
            date_str = datetime.now().strftime('%d/%m/%Y')
        except:
            time_str = datetime.now().strftime('%H:%M:%S')
            date_str = datetime.now().strftime('%d/%m/%Y')
        
        create_notification(
            f'🗑️ Item supprimé - "{item_name}" ({serial_number}) | {date_str} {time_str}',
            'success',
            serial_number,
            conn,
            cursor
        )
        
        conn.commit()
        conn.close()
        
        # Diffuser l'événement à tous les clients
        broadcast_event('items_changed', {'action': 'deleted', 'serialNumber': serial_number})
        broadcast_event('notifications_changed', {})
        
        return jsonify({'success': True}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

# ==================== API GROUPES/HIÉRARCHIE D'ITEMS ====================

@app.route('/api/items/<int:item_id>/set-parent', methods=['POST'])
def set_item_parent(item_id):
    """Définir un item comme enfant d'un autre item (créer une relation parent-enfant)"""
    try:
        data = request.get_json()
        parent_id = data.get('parentId')  # None pour retirer du groupe
        display_order = data.get('displayOrder', 0)
        
        conn = get_db()
        cursor = conn.cursor()
        
        # Vérifier que l'item existe
        cursor.execute('SELECT id FROM items WHERE id = ?', (item_id,))
        if not cursor.fetchone():
            conn.close()
            return jsonify({'success': False, 'error': 'Item non trouvé'}), 404
        
        # Si parent_id est fourni, vérifier qu'il existe et n'est pas l'item lui-même
        if parent_id is not None:
            if parent_id == item_id:
                conn.close()
                return jsonify({'success': False, 'error': 'Un item ne peut pas être son propre parent'}), 400
            
            cursor.execute('SELECT id FROM items WHERE id = ?', (parent_id,))
            if not cursor.fetchone():
                conn.close()
                return jsonify({'success': False, 'error': 'Item parent non trouvé'}), 404
            
            # Vérifier qu'on ne crée pas de boucle (l'item parent ne doit pas être un descendant de l'item actuel)
            cursor.execute('SELECT parent_id FROM items WHERE id = ?', (parent_id,))
            parent_row = cursor.fetchone()
            if parent_row and parent_row['parent_id'] == item_id:
                conn.close()
                return jsonify({'success': False, 'error': 'Impossible de créer une relation circulaire'}), 400
        
        # Mettre à jour la relation
        cursor.execute('''
            UPDATE items 
            SET parent_id = ?, display_order = ?
            WHERE id = ?
        ''', (parent_id, display_order, item_id))
        
        conn.commit()
        conn.close()
        
        # Diffuser l'événement
        broadcast_event('items_changed', {'action': 'hierarchy_updated', 'itemId': item_id})
        
        return jsonify({'success': True}), 200
        
    except Exception as e:
        print(f'[API] ERREUR POST /api/items/{item_id}/set-parent: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/items/<int:item_id>/remove-parent', methods=['POST'])
def remove_item_parent(item_id):
    """Retirer un item de son groupe (mettre parent_id à NULL)"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        cursor.execute('''
            UPDATE items 
            SET parent_id = NULL, display_order = 0
            WHERE id = ?
        ''', (item_id,))
        
        conn.commit()
        conn.close()
        
        # Diffuser l'événement
        broadcast_event('items_changed', {'action': 'hierarchy_updated', 'itemId': item_id})
        
        return jsonify({'success': True}), 200
        
    except Exception as e:
        print(f'[API] ERREUR POST /api/items/{item_id}/remove-parent: {str(e)}')
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/items/reorder-hierarchy', methods=['POST'])
def reorder_item_hierarchy():
    """Réorganiser l'ordre des items dans la hiérarchie"""
    try:
        data = request.get_json()
        items_order = data.get('items', [])  # [{id: 1, parentId: null, displayOrder: 0}, ...]
        
        conn = get_db()
        cursor = conn.cursor()
        
        # Mettre à jour chaque item
        for item_data in items_order:
            item_id = item_data.get('id')
            parent_id = item_data.get('parentId')
            display_order = item_data.get('displayOrder', 0)
            
            cursor.execute('''
                UPDATE items 
                SET parent_id = ?, display_order = ?
                WHERE id = ?
            ''', (parent_id, display_order, item_id))
        
        conn.commit()
        conn.close()
        
        # Diffuser l'événement
        broadcast_event('items_changed', {'action': 'hierarchy_reordered'})
        
        return jsonify({'success': True}), 200
        
    except Exception as e:
        print(f'[API] ERREUR POST /api/items/reorder-hierarchy: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

# ==================== API CATEGORIES ====================

@app.route('/api/categories', methods=['GET'])
def get_categories():
    """Récupérer toutes les catégories disponibles"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        cursor.execute('SELECT name FROM custom_categories')
        custom_categories = [row['name'] for row in cursor.fetchall()]
        
        cursor.execute('SELECT name FROM deleted_categories')
        deleted_categories = [row['name'] for row in cursor.fetchall()]
        
        conn.close()
        
        default_categories = [
            'materiel', 
            'drone', 
            'video', 
            'audio', 
            'streaming', 
            'robot', 
            'ordinateur',
            'casque_vr',
            'camera',
            'eclairage',
            'accessoire',
            'autre'
        ]
        available_default = [c for c in default_categories if c not in deleted_categories]
        available_custom = [c for c in custom_categories if c not in deleted_categories]
        
        return jsonify({
            'success': True,
            'categories': available_default + available_custom,
            'customCategories': custom_categories,
            'deletedCategories': deleted_categories
        }), 200
    except Exception as e:
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/categories', methods=['POST'])
def create_category():
    """Créer une nouvelle catégorie"""
    try:
        data = request.get_json()
        category_name = data.get('name', '').lower().strip()
        
        if not category_name:
            return jsonify({'success': False, 'error': 'Le nom de la catégorie est obligatoire'}), 400
        
        conn = get_db()
        cursor = conn.cursor()
        
        # Vérifier si la catégorie est supprimée (réactivation)
        cursor.execute('SELECT * FROM deleted_categories WHERE name = ?', (category_name,))
        if cursor.fetchone():
            cursor.execute('DELETE FROM deleted_categories WHERE name = ?', (category_name,))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Catégorie réactivée'}), 200
        
        # Créer la catégorie
        cursor.execute(
            'INSERT INTO custom_categories (name, created_at) VALUES (?, ?)',
            (category_name, datetime.now().isoformat())
        )
        
        # Créer une notification
        create_notification(f'Catégorie "{category_name}" ajoutée', 'success', None, conn, cursor)
        
        conn.commit()
        conn.close()
        
        # Diffuser l'événement à tous les clients
        broadcast_event('categories_changed', {'action': 'created', 'category': category_name})
        broadcast_event('notifications_changed', {})
        
        return jsonify({'success': True}), 201
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': 'Cette catégorie existe déjà'}), 409
    except Exception as e:
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/categories/<category_name>', methods=['DELETE'])
def delete_category(category_name):
    """Supprimer une catégorie"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM deleted_categories WHERE name = ?', (category_name,))
        if not cursor.fetchone():
            cursor.execute(
                'INSERT INTO deleted_categories (name, deleted_at) VALUES (?, ?)',
                (category_name, datetime.now().isoformat())
            )
        
        cursor.execute('DELETE FROM custom_categories WHERE name = ?', (category_name,))
        
        cursor.execute(
            'UPDATE items SET category = ?, last_updated = ? WHERE category = ?',
            ('autre', datetime.now().isoformat(), category_name)
        )
        
        updated_count = cursor.rowcount
        
        # Créer une notification
        create_notification(f'Catégorie "{category_name}" supprimée. {updated_count} item(s) mis à jour.', 'success', None, conn, cursor)
        
        conn.commit()
        conn.close()
        
        # Diffuser l'événement à tous les clients (les items ont été modifiés)
        broadcast_event('items_changed', {'action': 'category_deleted', 'category': category_name, 'updatedCount': updated_count})
        broadcast_event('categories_changed', {'action': 'deleted', 'category': category_name})
        broadcast_event('notifications_changed', {})
        
        return jsonify({
            'success': True,
            'message': f'Catégorie supprimée. {updated_count} item(s) mis à jour.'
        }), 200
    except Exception as e:
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

# ==================== API CUSTOM FIELDS (Colonnes personnalisées) ====================

@app.route('/api/custom-fields', methods=['GET'])
def get_custom_fields():
    """Récupérer tous les champs personnalisés"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, name, field_key, field_type, options, required, display_order, created_at
            FROM custom_fields
            ORDER BY display_order ASC, name ASC
        ''')
        
        rows = cursor.fetchall()
        fields = [{
            'id': row['id'],
            'name': row['name'],
            'fieldKey': row['field_key'],
            'fieldType': row['field_type'],
            'options': json.loads(row['options']) if row['options'] else None,
            'required': bool(row['required']),
            'displayOrder': row['display_order'],
            'createdAt': row['created_at']
        } for row in rows]
        
        conn.close()
        return jsonify({'success': True, 'fields': fields}), 200
    except Exception as e:
        print(f'[API] ERREUR GET /api/custom-fields: {str(e)}')
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/custom-fields', methods=['POST'])
def create_custom_field():
    """Créer un nouveau champ personnalisé"""
    try:
        data = request.get_json()
        name = data.get('name', '').strip()
        field_type = data.get('fieldType', 'text')
        options = data.get('options')  # Pour les champs de type 'select'
        required = data.get('required', False)
        
        if not name:
            return jsonify({'success': False, 'error': 'Le nom du champ est obligatoire'}), 400
        
        # Générer une clé unique à partir du nom
        field_key = re.sub(r'[^a-z0-9]', '_', name.lower())
        field_key = re.sub(r'_+', '_', field_key).strip('_')
        
        # Vérifier que le type est valide
        valid_types = ['text', 'number', 'date', 'select', 'checkbox', 'textarea', 'url', 'email']
        if field_type not in valid_types:
            return jsonify({'success': False, 'error': f'Type invalide. Types valides: {", ".join(valid_types)}'}), 400
        
        conn = get_db()
        cursor = conn.cursor()
        
        # Récupérer le prochain ordre d'affichage
        cursor.execute('SELECT MAX(display_order) FROM custom_fields')
        max_order = cursor.fetchone()[0] or 0
        
        cursor.execute('''
            INSERT INTO custom_fields (name, field_key, field_type, options, required, display_order, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (
            name,
            field_key,
            field_type,
            json.dumps(options) if options else None,
            1 if required else 0,
            max_order + 1,
            datetime.now().isoformat()
        ))
        
        field_id = cursor.lastrowid
        conn.commit()
        conn.close()
        
        print(f'[API] Champ personnalisé créé: {name} (type: {field_type})')
        
        # Diffuser l'événement
        broadcast_event('custom_fields_changed', {'action': 'created', 'fieldId': field_id, 'name': name})
        
        return jsonify({
            'success': True,
            'id': field_id,
            'fieldKey': field_key
        }), 201
        
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': 'Un champ avec ce nom existe déjà'}), 409
    except Exception as e:
        print(f'[API] ERREUR POST /api/custom-fields: {str(e)}')
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/custom-fields/<int:field_id>', methods=['PUT'])
def update_custom_field(field_id):
    """Mettre à jour un champ personnalisé"""
    try:
        data = request.get_json()
        name = data.get('name', '').strip()
        field_type = data.get('fieldType')
        options = data.get('options')
        required = data.get('required')
        display_order = data.get('displayOrder')
        
        conn = get_db()
        cursor = conn.cursor()
        
        # Construire la requête de mise à jour dynamiquement
        updates = []
        params = []
        
        if name:
            updates.append('name = ?')
            params.append(name)
            # Mettre à jour aussi la clé
            field_key = re.sub(r'[^a-z0-9]', '_', name.lower())
            field_key = re.sub(r'_+', '_', field_key).strip('_')
            updates.append('field_key = ?')
            params.append(field_key)
        
        if field_type:
            updates.append('field_type = ?')
            params.append(field_type)
        
        if options is not None:
            updates.append('options = ?')
            params.append(json.dumps(options) if options else None)
        
        if required is not None:
            updates.append('required = ?')
            params.append(1 if required else 0)
        
        if display_order is not None:
            updates.append('display_order = ?')
            params.append(display_order)
        
        if not updates:
            return jsonify({'success': False, 'error': 'Aucune donnée à mettre à jour'}), 400
        
        params.append(field_id)
        cursor.execute(f'''
            UPDATE custom_fields
            SET {', '.join(updates)}
            WHERE id = ?
        ''', params)
        
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({'success': False, 'error': 'Champ non trouvé'}), 404
        
        conn.commit()
        conn.close()
        
        # Diffuser l'événement
        broadcast_event('custom_fields_changed', {'action': 'updated', 'fieldId': field_id})
        
        return jsonify({'success': True}), 200
        
    except Exception as e:
        print(f'[API] ERREUR PUT /api/custom-fields/{field_id}: {str(e)}')
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/custom-fields/<int:field_id>', methods=['DELETE'])
def delete_custom_field(field_id):
    """Supprimer un champ personnalisé"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        # Récupérer le nom du champ avant suppression
        cursor.execute('SELECT name, field_key FROM custom_fields WHERE id = ?', (field_id,))
        field = cursor.fetchone()
        
        if not field:
            conn.close()
            return jsonify({'success': False, 'error': 'Champ non trouvé'}), 404
        
        field_name = field['name']
        field_key = field['field_key']
        
        # Supprimer le champ de la table custom_fields
        cursor.execute('DELETE FROM custom_fields WHERE id = ?', (field_id,))
        
        # Optionnel: Supprimer les données de ce champ dans tous les items
        # (on garde les données pour l'instant, au cas où)
        
        conn.commit()
        conn.close()
        
        print(f'[API] Champ personnalisé supprimé: {field_name}')
        
        # Diffuser l'événement
        broadcast_event('custom_fields_changed', {'action': 'deleted', 'fieldId': field_id, 'fieldKey': field_key})
        
        return jsonify({'success': True}), 200
        
    except Exception as e:
        print(f'[API] ERREUR DELETE /api/custom-fields/{field_id}: {str(e)}')
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

# ==================== API NOTIFICATIONS ====================

@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    """Récupérer les notifications"""
    try:
        print('[API] GET /api/notifications - Récupération des notifications...')
        conn = get_db()
        cursor = conn.cursor()
        
        # Récupérer les 50 dernières notifications (avec item_hex_id pour navigation)
        cursor.execute('''
            SELECT id, message, type, item_serial_number, item_hex_id, created_at
            FROM notifications
            ORDER BY created_at DESC
            LIMIT 50
        ''')
        
        rows = cursor.fetchall()
        notifications = []
        
        # Traiter chaque notification individuellement pour gérer les erreurs d'encodage
        for row in rows:
            row_dict = dict(row)  # sqlite3.Row n'a pas .get(), convertir en dict
            try:
                # Récupérer le message et le sanitizer
                raw_message = row_dict.get('message')
                if raw_message is None:
                    raw_message = ''
                
                # Sanitizer le message pour éviter les problèmes d'encodage
                clean_message = sanitize_notification_message(raw_message)
                
                notifications.append({
                    'id': row_dict.get('id'),
                    'message': clean_message,
                    'type': row_dict.get('type'),
                    'itemSerialNumber': row_dict.get('item_serial_number'),
                    'itemHexId': row_dict.get('item_hex_id'),
                    'timestamp': row_dict.get('created_at'),
                    'created_at': row_dict.get('created_at')  # Alias pour compatibilité
                })
            except Exception as msg_error:
                # Si une notification spécifique cause une erreur, la remplacer par un message par défaut
                try:
                    error_msg = str(msg_error)
                    # S'assurer que le message peut être encodé en UTF-8
                    error_msg = error_msg.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
                    safe_error_msg = sanitize_error(error_msg)
                    print(f'[API] Erreur lors du traitement d\'une notification (ID: {row_dict.get("id", "unknown")}): {safe_error_msg}')
                except:
                    print(f'[API] Erreur lors du traitement d\'une notification (ID: {row_dict.get("id", "unknown")})')
                notifications.append({
                    'id': row_dict.get('id', 0),
                    'message': 'Message de notification (erreur d\'encodage)',
                    'type': row_dict.get('type', 'info'),
                    'itemSerialNumber': row_dict.get('item_serial_number'),
                    'itemHexId': row_dict.get('item_hex_id'),
                    'timestamp': row_dict.get('created_at', ''),
                    'created_at': row_dict.get('created_at', '')
                })
        
        conn.close()
        print(f'[API] GET /api/notifications - {len(notifications)} notifications retournées')
        # Ne pas afficher les messages dans la console car ils peuvent contenir des caractères Unicode
        # qui causent des erreurs d'encodage sur Windows
        return jsonify({'success': True, 'notifications': notifications}), 200
    except Exception as e:
        # Gérer les erreurs d'encodage de manière plus robuste
        error_msg = str(e)
        # S'assurer que le message d'erreur peut être encodé en UTF-8
        try:
            # Encoder en UTF-8 avec remplacement des caractères problématiques
            error_msg = error_msg.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
        except:
            error_msg = 'Erreur lors de la récupération des notifications'
        
        # Utiliser sanitize_error pour nettoyer le message avant l'affichage
        safe_error_msg = sanitize_error(error_msg)
        try:
            print(f'[API] ERREUR GET /api/notifications: {safe_error_msg}')
        except:
            # Si l'affichage échoue encore, utiliser un message générique
            print('[API] ERREUR GET /api/notifications: Erreur lors de la récupération des notifications')
        
        import traceback
        try:
            # Capturer la traceback dans une chaîne pour éviter les problèmes d'encodage
            import io
            traceback_buffer = io.StringIO()
            traceback.print_exc(file=traceback_buffer)
            traceback_str = traceback_buffer.getvalue()
            # Nettoyer la traceback avant l'affichage
            traceback_str = traceback_str.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
            print(traceback_str)
        except:
            # Si même l'affichage de la traceback échoue, ignorer
            pass
        
        return jsonify({'success': False, 'error': safe_error_msg}), 500

@app.route('/api/notifications/<int:notification_id>', methods=['DELETE'])
def delete_notification(notification_id):
    """Supprimer une notification spécifique"""
    try:
        print(f'[API] DELETE /api/notifications/{notification_id} - Suppression de la notification...')
        conn = get_db()
        cursor = conn.cursor()
        
        # Vérifier que la notification existe
        cursor.execute('SELECT id FROM notifications WHERE id = ?', (notification_id,))
        if not cursor.fetchone():
            conn.close()
            print(f'[API] DELETE /api/notifications/{notification_id} - Notification non trouvée')
            return jsonify({'success': False, 'error': 'Notification non trouvée'}), 404
        
        # Supprimer la notification
        cursor.execute('DELETE FROM notifications WHERE id = ?', (notification_id,))
        
        conn.commit()
        conn.close()
        
        print(f'[API] DELETE /api/notifications/{notification_id} - Notification supprimée avec succès')
        
        # Diffuser l'événement
        broadcast_event('notifications_changed', {})
        
        return jsonify({'success': True}), 200
    except Exception as e:
        print(f'[API] ERREUR DELETE /api/notifications/{notification_id}: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/notifications', methods=['DELETE'])
def clear_notifications():
    """Effacer toutes les notifications"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        cursor.execute('DELETE FROM notifications')
        
        conn.commit()
        conn.close()
        
        # Diffuser l'événement
        broadcast_event('notifications_changed', {})
        
        return jsonify({'success': True}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

# ==================== PROXY POUR APIs EXTERNES ====================

@app.route('/api/proxy/gtinsearch', methods=['GET'])
def proxy_gtinsearch():
    """Proxy pour recherche de produit par code-barres (avec fallback multi-APIs)"""
    try:
        gtin = request.args.get('gtin')
        if not gtin:
            _log('ERREUR', 'Paramètre gtin manquant')
            return jsonify({'success': False, 'error': 'Paramètre gtin manquant'}), 400
        
        _log('INFO', f'Recherche produit pour code: {gtin}')
        
        # Headers communs pour simuler un navigateur
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36',
            'Accept': 'application/json, text/plain, */*',
            'Accept-Language': 'fr-FR,fr;q=0.9,en-US;q=0.8,en;q=0.7',
            'Referer': 'https://www.google.com/',
        }
        
        # 1. Essayer Open Food Facts d'abord (gratuit, pas de limite)
        try:
            _log('INFO', 'Essai Open Food Facts...')
            off_url = f'https://world.openfoodfacts.org/api/v0/product/{gtin}.json'
            off_response = requests.get(off_url, timeout=8, headers=headers)
            
            if off_response.status_code == 200:
                off_data = off_response.json()
                if off_data.get('status') == 1 and off_data.get('product'):
                    product = off_data['product']
                    name = product.get('product_name') or product.get('product_name_fr') or product.get('generic_name')
                    if name:
                        _log('INFO', f'Open Food Facts: trouvé "{name}"')
                        # Récupérer les images disponibles
                        images = []
                        if product.get('image_url'):
                            images.append(product.get('image_url'))
                        elif product.get('image_front_url'):
                            images.append(product.get('image_front_url'))
                        
                        return jsonify({
                            'success': True,
                            'name': name,
                            'brand': product.get('brands', ''),
                            'category': product.get('categories', ''),
                            'image': images[0] if images else '',
                            'images': images,
                            'source': 'Open Food Facts'
                        }), 200
        except Exception as e:
            _log('INFO', f'Open Food Facts erreur: {str(e)}')
        
        # 2. Essayer UPC Item DB (gratuit, limité)
        try:
            _log('INFO', 'Essai UPC Item DB...')
            upc_url = f'https://api.upcitemdb.com/prod/trial/lookup?upc={gtin}'
            upc_response = requests.get(upc_url, timeout=8, headers={
                **headers,
                'Accept': 'application/json',
            })
            
            if upc_response.status_code == 200:
                upc_data = upc_response.json()
                if upc_data.get('items') and len(upc_data['items']) > 0:
                    item = upc_data['items'][0]
                    name = item.get('title')
                    if name:
                        _log('INFO', f'UPC Item DB: trouvé "{name}"')
                        images = item.get('images', [])
                        return jsonify({
                            'success': True,
                            'name': name,
                            'brand': item.get('brand', ''),
                            'category': item.get('category', ''),
                            'image': images[0] if images else '',
                            'images': images,
                            'source': 'UPC Item DB'
                        }), 200
        except Exception as e:
            _log('INFO', f'UPC Item DB erreur: {str(e)}')
        
        # 3. Essayer GTINsearch en dernier (souvent bloqué)
        try:
            _log('INFO', 'Essai GTINsearch...')
            gtin_url = f'https://gtinsearch.org/api?gtin={gtin}'
            gtin_response = requests.get(gtin_url, timeout=8, headers=headers)
            
            if gtin_response.status_code == 200:
                gtin_data = gtin_response.json()
                if gtin_data.get('name'):
                    _log('INFO', f'GTINsearch: trouvé "{gtin_data["name"]}"')
                    return jsonify({
                        'success': True,
                        'name': gtin_data['name'],
                        'brand': gtin_data.get('brand', ''),
                        'category': gtin_data.get('category', ''),
                        'images': [],
                        'source': 'GTINsearch'
                    }), 200
            else:
                _log('INFO', f'GTINsearch: HTTP {gtin_response.status_code}')
        except Exception as e:
            _log('INFO', f'GTINsearch erreur: {str(e)}')
        
        # Aucune API n'a trouvé le produit
        _log('INFO', f'Aucun résultat trouvé pour: {gtin}')
        return jsonify({'success': False, 'error': 'Produit non trouvé', 'name': None}), 200
            
    except Exception as e:
        _log('ERREUR', f'Erreur proxy gtinsearch: {str(e)}', e)
        # S'assurer de retourner un JSON valide même en cas d'erreur
        try:
            return jsonify({'success': False, 'error': str(e)}), 200
        except:
            return jsonify({'success': False, 'error': 'Erreur interne'}), 200

@app.route('/api/proxy/openfoodfacts', methods=['GET'])
def proxy_openfoodfacts():
    """Proxy pour Open Food Facts (contourne CORS)"""
    try:
        barcode = request.args.get('barcode')
        if not barcode:
            return jsonify({'success': False, 'error': 'Paramètre barcode manquant'}), 400
        
        print(f'[Proxy] Requête Open Food Facts pour: {barcode}')
        url = f'https://world.openfoodfacts.org/api/v0/product/{barcode}.json'
        
        response = requests.get(url, timeout=10, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        
        print(f'[Proxy] Open Food Facts status: {response.status_code}')
        
        if response.status_code != 200:
            print(f'[Proxy] Open Food Facts erreur HTTP: {response.status_code}')
            return jsonify({'success': False, 'status': 0, 'error': f'HTTP {response.status_code}'}), 200
        
        try:
            data = response.json()
            print(f'[Proxy] Open Food Facts réponse reçue')
            return jsonify(data), 200
        except ValueError:
            print(f'[Proxy] Open Food Facts réponse non-JSON')
            return jsonify({'success': False, 'status': 0, 'error': 'Réponse invalide'}), 200
            
    except requests.exceptions.Timeout:
        print(f'[Proxy] Open Food Facts timeout')
        return jsonify({'success': False, 'status': 0, 'error': 'Timeout'}), 200
    except requests.exceptions.RequestException as e:
        print(f'[Proxy] Open Food Facts erreur: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'status': 0, 'error': sanitize_error(e)}), 200
    except Exception as e:
        print(f'[Proxy] Open Food Facts erreur inattendue: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'status': 0, 'error': sanitize_error(e)}), 200

@app.route('/api/proxy/openfoodfacts/search', methods=['GET'])
def proxy_openfoodfacts_search():
    """Proxy pour la recherche Open Food Facts (contourne CORS)"""
    try:
        query = request.args.get('query')
        if not query:
            return jsonify({'success': False, 'error': 'Paramètre query manquant'}), 400
        
        print(f'[Proxy] Recherche Open Food Facts pour: {query}')
        url = f'https://world.openfoodfacts.org/cgi/search.pl?search_terms={query}&search_simple=1&action=process&json=1&page_size=8'
        
        response = requests.get(url, timeout=10, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        
        print(f'[Proxy] Open Food Facts recherche status: {response.status_code}')
        
        if response.status_code != 200:
            print(f'[Proxy] Open Food Facts recherche erreur HTTP: {response.status_code}')
            return jsonify({'success': False, 'products': [], 'error': f'HTTP {response.status_code}'}), 200
        
        try:
            data = response.json()
            print(f'[Proxy] Open Food Facts recherche réponse reçue')
            return jsonify(data), 200
        except ValueError:
            print(f'[Proxy] Open Food Facts recherche réponse non-JSON')
            return jsonify({'success': False, 'products': [], 'error': 'Réponse invalide'}), 200
            
    except requests.exceptions.Timeout:
        print(f'[Proxy] Open Food Facts recherche timeout')
        return jsonify({'success': False, 'products': [], 'error': 'Timeout'}), 200
    except requests.exceptions.RequestException as e:
        print(f'[Proxy] Open Food Facts recherche erreur réseau: {str(e)}')
        return jsonify({'success': False, 'products': [], 'error': str(e)}), 200
    except Exception as e:
        print(f'[Proxy] Open Food Facts recherche erreur inattendue: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'products': [], 'error': sanitize_error(e)}), 200

@app.route('/api/proxy/image', methods=['GET'])
def proxy_image():
    """Proxy pour récupérer des images externes et les convertir en base64 (contourne CORS)"""
    try:
        image_url = request.args.get('url')
        if not image_url:
            return jsonify({'success': False, 'error': 'Paramètre url manquant'}), 400
        
        print(f'[Proxy] Récupération image depuis: {image_url}')
        
        # Vérifier que l'URL est une image
        if not any(ext in image_url.lower() for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp']):
            print(f'[Proxy] URL ne semble pas être une image')
        
        # Récupérer l'image
        response = requests.get(image_url, timeout=15, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'image/*'
        })
        
        print(f'[Proxy] Image status: {response.status_code}')
        
        if response.status_code != 200:
            print(f'[Proxy] Erreur HTTP lors de la récupération: {response.status_code}')
            return jsonify({'success': False, 'error': f'HTTP {response.status_code}'}), 200
        
        # Vérifier le content-type
        content_type = response.headers.get('Content-Type', '')
        print(f'[Proxy] Content-Type: {content_type}')
        
        if not content_type.startswith('image/'):
            print(f'[Proxy] Content-Type invalide: {content_type}')
            return jsonify({'success': False, 'error': f'Type de contenu invalide: {content_type}'}), 200
        
        # Convertir en base64
        import base64
        image_base64 = base64.b64encode(response.content).decode('utf-8')
        data_uri = f'data:{content_type};base64,{image_base64}'
        
        print(f'[Proxy] Image converted to base64 ({len(image_base64)} chars)')
        
        return jsonify({
            'success': True,
            'image': data_uri,
            'contentType': content_type
        }), 200
        
    except requests.exceptions.Timeout:
        print(f'[Proxy] Timeout lors de la récupération de l\'image')
        return jsonify({'success': False, 'error': 'Timeout'}), 200
    except requests.exceptions.RequestException as e:
        print(f'[Proxy] Erreur réseau: {str(e)}')
        return jsonify({'success': False, 'error': f'Erreur réseau: {str(e)}'}), 200
    except Exception as e:
        print(f'[Proxy] Erreur inattendue: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 200
            
    except requests.exceptions.Timeout:
        print(f'[Proxy] Open Food Facts recherche timeout')
        return jsonify({'success': False, 'products': [], 'error': 'Timeout'}), 200
    except requests.exceptions.RequestException as e:
        print(f'[Proxy] Open Food Facts recherche erreur: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'products': [], 'error': sanitize_error(e)}), 200
    except Exception as e:
        print(f'[Proxy] Open Food Facts recherche erreur inattendue: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'products': [], 'error': sanitize_error(e)}), 200

# ==================== API LOCATIONS ====================

@app.route('/api/rentals', methods=['GET'])
def get_rentals():
    """Récupérer toutes les locations"""
    try:
        status_filter = request.args.get('status', '')
        conn = get_db()
        cursor = conn.cursor()
        
        if status_filter:
            cursor.execute('''
                SELECT * FROM rentals
                WHERE status = ?
                ORDER BY start_date DESC
            ''', (status_filter,))
        else:
            cursor.execute('''
                SELECT * FROM rentals
                ORDER BY start_date DESC
            ''')
        
        rows = cursor.fetchall()
        rentals = [{
            'id': row['id'],
            'renterName': row['renter_name'],
            'renterEmail': row['renter_email'],
            'renterPhone': row['renter_phone'],
            'renterAddress': row['renter_address'],
            'rentalPrice': row['rental_price'],
            'rentalDeposit': row['rental_deposit'],
            'rentalDuration': row['rental_duration'],
            'startDate': row['start_date'],
            'endDate': row['end_date'],
            'status': row['status'],
            'itemsData': json.loads(row['items_data']),
            'createdAt': row['created_at'],
            'updatedAt': row['updated_at']
        } for row in rows]
        
        conn.close()
        return jsonify({'success': True, 'rentals': rentals}), 200
    except Exception as e:
        print(f'[API] ERREUR GET /api/rentals: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/rentals', methods=['POST'])
def create_rental():
    """Créer une nouvelle location"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Données JSON invalides'}), 400
        
        # Validation des champs requis (email et téléphone sont optionnels)
        required_fields = ['renterName', 'rentalPrice', 'rentalDeposit', 'rentalDuration', 'startDate', 'endDate', 'itemsData']
        missing_fields = validate_required_fields(data, required_fields)
        if missing_fields:
            return jsonify({'success': False, 'error': f'Champs obligatoires manquants: {", ".join(missing_fields)}'}), 400
        
        # Validation de l'email (optionnel, mais doit être valide si fourni)
        renter_email = data.get('renterEmail', '').strip()
        if renter_email and not validate_email(renter_email):
            return jsonify({'success': False, 'error': f'Format d\'email invalide: {renter_email}'}), 400
        
        # Validation du téléphone (optionnel, mais doit être valide si fourni)
        renter_phone = data.get('renterPhone', '').strip()
        if renter_phone and not validate_phone(renter_phone):
            return jsonify({'success': False, 'error': f'Format de téléphone invalide: {renter_phone}'}), 400
        
        # Validation des montants
        if not validate_positive_number(data.get('rentalPrice')):
            return jsonify({'success': False, 'error': 'Le prix de location doit être un nombre positif'}), 400
        if not validate_positive_number(data.get('rentalDeposit')):
            return jsonify({'success': False, 'error': 'La caution doit être un nombre positif'}), 400
        
        # Validation de la durée
        if not validate_positive_number(data.get('rentalDuration'), allow_zero=False):
            return jsonify({'success': False, 'error': 'La durée doit être un nombre positif'}), 400
        
        # Validation des items
        if not isinstance(data.get('itemsData'), list) or len(data.get('itemsData', [])) == 0:
            return jsonify({'success': False, 'error': 'Au moins un item doit être sélectionné'}), 400
        
        # Sanitization
        data['renterName'] = sanitize_string(data.get('renterName'), 200)
        data['renterEmail'] = sanitize_string(data.get('renterEmail'), 200)
        data['renterPhone'] = sanitize_string(data.get('renterPhone'), 50)
        data['renterAddress'] = sanitize_string(data.get('renterAddress'), 500)
        data['notes'] = sanitize_string(data.get('notes'), 2000)
        
        now = datetime.now().isoformat()
        conn = get_db()
        cursor = conn.cursor()
        
        cursor.execute('''
            INSERT INTO rentals (
                renter_name, renter_email, renter_phone, renter_address,
                rental_price, rental_deposit, rental_duration,
                start_date, end_date, status, items_data,
                created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            data['renterName'],
            data['renterEmail'],
            data['renterPhone'],
            data.get('renterAddress', ''),
            data['rentalPrice'],
            data['rentalDeposit'],
            data['rentalDuration'],
            data['startDate'],
            data['endDate'],
            data.get('status', 'en_cours'),
            json.dumps(data['itemsData']),
            now,
            now
        ))
        
        rental_id = cursor.lastrowid
        
        # Mettre à jour le statut des items dans l'inventaire
        # Déterminer le statut selon la date de début
        item_status = 'location_future' if data.get('status') == 'a_venir' else 'en_location'
        
        for item_data in data['itemsData']:
            serial_number = item_data.get('serialNumber')
            quantity = item_data.get('quantity', 1)  # Quantité louée
            
            if serial_number:
                try:
                    # Récupérer l'item actuel
                    cursor.execute('SELECT quantity FROM items WHERE serial_number = ?', (serial_number,))
                    row = cursor.fetchone()
                    
                    if row:
                        current_quantity = row['quantity'] or 1
                        remaining_quantity = max(0, current_quantity - quantity)
                        
                        # Mettre à jour l'item
                        cursor.execute('''
                            UPDATE items 
                            SET status = ?,
                                quantity = ?,
                                rental_end_date = ?,
                                current_rental_id = ?,
                                last_updated = ?
                            WHERE serial_number = ?
                        ''', (
                            item_status if remaining_quantity == 0 else 'en_stock',  # Si tout est loué, changer le statut
                            remaining_quantity,
                            data['endDate'],
                            rental_id,
                            datetime.now().isoformat(),
                            serial_number
                        ))
                        
                        print(f'[RENTAL] Item {serial_number}: {quantity}/{current_quantity} loués, {remaining_quantity} restants, statut: {item_status if remaining_quantity == 0 else "en_stock"}')
                except Exception as e:
                    print(f'[RENTAL] Erreur mise à jour item {serial_number}: {e}')
        
        conn.commit()
        conn.close()
        
        # Diffuser les événements
        broadcast_event('rentals_changed', {'action': 'created', 'id': rental_id})
        broadcast_event('items_changed', {'action': 'updated', 'rental_id': rental_id})
        
        return jsonify({'success': True, 'id': rental_id}), 201
    except Exception as e:
        print(f'[API] ERREUR POST /api/rentals: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/rentals/<int:rental_id>', methods=['PUT'])
def update_rental(rental_id):
    """Mettre à jour une location"""
    try:
        data = request.get_json()
        
        # Validation de l'email (optionnel, mais doit être valide si fourni)
        renter_email = data.get('renterEmail', '').strip()
        if renter_email and not validate_email(renter_email):
            return jsonify({'success': False, 'error': f'Format d\'email invalide: {renter_email}'}), 400
        
        # Validation du téléphone (optionnel, mais doit être valide si fourni)
        renter_phone = data.get('renterPhone', '').strip()
        if renter_phone and not validate_phone(renter_phone):
            return jsonify({'success': False, 'error': f'Format de téléphone invalide: {renter_phone}'}), 400
        
        now = datetime.now().isoformat()
        conn = get_db()
        cursor = conn.cursor()
        
        # Récupérer l'ancien statut avant mise à jour
        cursor.execute('SELECT status, items_data FROM rentals WHERE id = ?', (rental_id,))
        old_rental = cursor.fetchone()
        old_status = old_rental['status'] if old_rental else None
        
        cursor.execute('''
            UPDATE rentals
            SET renter_name = ?, renter_email = ?, renter_phone = ?, renter_address = ?,
                rental_price = ?, rental_deposit = ?, rental_duration = ?,
                start_date = ?, end_date = ?, status = ?, items_data = ?,
                updated_at = ?
            WHERE id = ?
        ''', (
            data['renterName'],
            renter_email,
            renter_phone,
            data.get('renterAddress', ''),
            data['rentalPrice'],
            data['rentalDeposit'],
            data['rentalDuration'],
            data['startDate'],
            data['endDate'],
            data['status'],
            json.dumps(data['itemsData']),
            now,
            rental_id
        ))
        
        # Si le statut passe à 'termine', libérer les items
        if old_status and old_status != 'termine' and data['status'] == 'termine':
            if data.get('itemsData'):
                for item_data in data['itemsData']:
                    serial_number = item_data.get('serialNumber')
                    quantity = item_data.get('quantity', 1)
                    
                    if serial_number:
                        # Récupérer la quantité actuelle
                        cursor.execute('SELECT quantity FROM items WHERE serial_number = ?', (serial_number,))
                        item_row = cursor.fetchone()
                        
                        if item_row:
                            current_quantity = item_row['quantity'] or 0
                            new_quantity = current_quantity + quantity
                            
                            # Libérer les quantités
                            cursor.execute('''
                                UPDATE items 
                                SET status = 'en_stock',
                                    quantity = ?,
                                    rental_end_date = NULL,
                                    current_rental_id = NULL,
                                    last_updated = ?
                                WHERE serial_number = ?
                            ''', (
                                new_quantity,
                                now,
                                serial_number
                            ))
                            
                            print(f'[RENTAL UPDATE] Location terminée - Item {serial_number}: {quantity} libérés, nouveau total: {new_quantity}')
        
        # Si le statut passe de 'a_venir' à 'en_cours', mettre à jour le statut des items
        elif old_status == 'a_venir' and data['status'] == 'en_cours':
            if data.get('itemsData'):
                for item_data in data['itemsData']:
                    serial_number = item_data.get('serialNumber')
                    if serial_number:
                        cursor.execute('''
                            UPDATE items 
                            SET status = 'en_location'
                            WHERE serial_number = ? AND current_rental_id = ?
                        ''', (serial_number, rental_id))
                        print(f'[RENTAL UPDATE] Statut changé: location_future -> en_location pour {serial_number}')
        
        conn.commit()
        conn.close()
        
        # Diffuser les événements
        broadcast_event('rentals_changed', {'action': 'updated', 'id': rental_id})
        broadcast_event('items_changed', {'action': 'updated', 'rental_id': rental_id})
        
        return jsonify({'success': True}), 200
    except Exception as e:
        print(f'[API] ERREUR PUT /api/rentals/{rental_id}: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/rentals/<int:rental_id>', methods=['DELETE'])
def delete_rental(rental_id):
    """Supprimer une location et libérer les items"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        # Récupérer les items de la location avant de la supprimer
        cursor.execute('SELECT items_data FROM rentals WHERE id = ?', (rental_id,))
        row = cursor.fetchone()
        
        if row and row['items_data']:
            try:
                items_data = json.loads(row['items_data'])
                
                # Libérer les quantités dans l'inventaire
                for item_data in items_data:
                    serial_number = item_data.get('serialNumber')
                    quantity = item_data.get('quantity', 1)
                    
                    if serial_number:
                        # Récupérer la quantité actuelle
                        cursor.execute('SELECT quantity FROM items WHERE serial_number = ?', (serial_number,))
                        item_row = cursor.fetchone()
                        
                        if item_row:
                            current_quantity = item_row['quantity'] or 0
                            new_quantity = current_quantity + quantity
                            
                            # Libérer les quantités et remettre le statut à en_stock
                            cursor.execute('''
                                UPDATE items 
                                SET status = 'en_stock',
                                    quantity = ?,
                                    rental_end_date = NULL,
                                    current_rental_id = NULL,
                                    last_updated = ?
                                WHERE serial_number = ?
                            ''', (
                                new_quantity,
                                datetime.now().isoformat(),
                                serial_number
                            ))
                            
                            print(f'[RENTAL DELETE] Item {serial_number}: {quantity} libérés, nouveau total: {new_quantity}')
            except Exception as e:
                print(f'[RENTAL DELETE] Erreur libération items: {e}')
        
        cursor.execute('DELETE FROM rentals WHERE id = ?', (rental_id,))
        conn.commit()
        conn.close()
        
        # Diffuser l'événement
        broadcast_event('rentals_changed', {'action': 'deleted', 'id': rental_id})
        
        return jsonify({'success': True}), 200
    except Exception as e:
        print(f'[API] ERREUR DELETE /api/rentals/{rental_id}: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

def generate_rental_caution_pdf(rental):
    """Générer un PDF de caution à partir des données location avec fpdf."""
    # Utiliser fpdf directement (basé sur pdf.py)
    if FPDF_AVAILABLE:
        return generate_rental_caution_from_template(rental, None)
    
    # Sinon, fallback vers la génération basique avec reportlab
    if not PDF_AVAILABLE:
        return None
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name='CustomTitle', parent=styles['Heading1'], fontSize=16, spaceAfter=12)
    heading_style = ParagraphStyle(name='CustomHeading', parent=styles['Heading2'], fontSize=12, spaceAfter=6)
    normal_style = styles['Normal']
    elements = []
    
    def format_date(date_str):
        if not date_str:
            return ''
        try:
            date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
            return date_obj.strftime('%d/%m/%Y')
        except:
            return date_str
    
    renter_name = rental.get('renter_name') or ''
    items_data = []
    try:
        items_data = json.loads(rental['items_data']) if rental.get('items_data') else []
    except:
        pass
    items_list_str = ', '.join([
        f"{item.get('name', 'Item')} ({item.get('brand', '')} {item.get('model', '')})"
        for item in items_data
    ]) if items_data else '—'
    
    elements.append(Paragraph('Modèle pour caution Location', title_style))
    elements.append(Spacer(1, 12))
    
    elements.append(Paragraph('Informations du locataire', heading_style))
    locataire_data = [
        ['Nom complet :', renter_name],
        ['Adresse :', (rental.get('renter_address') or '—')],
        ['Email :', (rental.get('renter_email') or '—')],
        ['Téléphone :', (rental.get('renter_phone') or '—')],
    ]
    t1 = Table(locataire_data, colWidths=[4*cm, 12*cm])
    t1.setStyle(TableStyle([('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'), ('FONTSIZE', (0, 0), (-1, -1), 10)]))
    elements.append(t1)
    elements.append(Spacer(1, 14))
    
    elements.append(Paragraph('Période de location', heading_style))
    periode_data = [
        ['Date de début :', format_date(rental.get('start_date'))],
        ['Date de fin :', format_date(rental.get('end_date'))],
        ['Durée :', f"{rental.get('rental_duration') or 0} jour(s)"],
    ]
    t2 = Table(periode_data, colWidths=[4*cm, 12*cm])
    t2.setStyle(TableStyle([('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'), ('FONTSIZE', (0, 0), (-1, -1), 10)]))
    elements.append(t2)
    elements.append(Spacer(1, 14))
    
    elements.append(Paragraph('Montants', heading_style))
    montant_caution = rental.get('rental_deposit') or 0
    montant_location = rental.get('rental_price') or 0
    montants_data = [
        ['Montant location :', f"{montant_location:.2f} €"],
        ['Caution :', f"{montant_caution:.2f} €"],
    ]
    t3 = Table(montants_data, colWidths=[4*cm, 12*cm])
    t3.setStyle(TableStyle([('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'), ('FONTSIZE', (0, 0), (-1, -1), 10)]))
    elements.append(t3)
    elements.append(Spacer(1, 14))
    
    elements.append(Paragraph('Matériel loué', heading_style))
    elements.append(Paragraph(items_list_str, normal_style))
    elements.append(Spacer(1, 14))
    
    elements.append(Paragraph(f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}", ParagraphStyle(name='Small', parent=normal_style, fontSize=8, textColor=colors.gray)))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_rental_caution_from_template(rental, template_path):
    """Générer un PDF avec fpdf (basé sur pdf.py) avec les infos du loueur."""
    try:
        if not FPDF_AVAILABLE:
            print('[PDF] fpdf non disponible, impossible de générer le PDF')
            return None
        
        def format_date(date_str):
            if not date_str:
                return ''
            try:
                date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                return date_obj.strftime('%d/%m/%Y')
            except:
                return date_str
        
        # Préparer les données
        date_doc = datetime.now().strftime('%d/%m/%Y')
        nom_client = rental.get('renter_name') or 'Client'
        
        # Récupérer les items pour le matériel
        items_data = []
        try:
            items_data = json.loads(rental['items_data']) if rental.get('items_data') else []
        except:
            pass
        
        # Construire la liste du matériel
        if items_data:
            materiel = ', '.join([
                f"{item.get('name', '')} {item.get('brand', '')} {item.get('model', '')}".strip()
                for item in items_data
            ])
        else:
            materiel = 'matériel'
        
        date_retrait = format_date(rental.get('end_date'))
        montant_caution = f"{rental.get('rental_deposit') or 0:.2f} CHF"
        
        # --- Création du PDF avec fpdf (comme dans pdf.py) ---
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=9)
        
        # --- Contenu du document (template exact de pdf.py) ---
        texte = f"""
Date : {date_doc}

Encaissement de la Caution pour {materiel}

Entre GlobalVision Communication Sàrl et {nom_client}


Cher client,

Nous vous remercions d'avoir choisi GlobalVision pour la location de {materiel}.
Afin d'assurer la sécurité du matériel, nous vous demandons de fournir une caution
lors du retrait du kit, qui aura lieu le {date_retrait}.

Détails de la Caution :

Montant : {montant_caution}
Mode de Paiement : Espèces ou paiement par carte.
Si paiement par carte, un frais de 3% est appliqué lors du remboursement de la caution.

Conditions de Remboursement :
La caution sera intégralement restituée lors du dépôt du matériel à condition que
celui-ci soit rendu dans le même état que lors de son retrait. Une vérification sera
effectuée pour s'assurer que tous les éléments du kit sont présents et en bon état
de fonctionnement.

Nous vous remercions de votre confiance et de votre coopération.

Cordialement,

Global Vision Communication SARL

Date et Signature des deux parties :
"""
        
        pdf.multi_cell(0, 8, texte)
        
        # Sauvegarder dans un buffer au lieu d'un fichier
        buffer = BytesIO()
        pdf_output = pdf.output(dest='S').encode('latin-1')  # fpdf retourne une string
        buffer.write(pdf_output)
        buffer.seek(0)
        
        print('[PDF] PDF généré avec succès avec fpdf')
        return buffer
        
    except Exception as e:
        print(f'[PDF] Erreur lors de la génération avec fpdf: {str(e)}')
        import traceback
        safe_traceback()
        return None


@app.route('/api/rentals/<int:rental_id>/caution-doc', methods=['GET'])
def get_rental_caution_doc(rental_id):
    """Générer et télécharger le document de caution (PDF par défaut, DOCX si format=docx)"""
    try:
        # Récupérer la location
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM rentals WHERE id = ?', (rental_id,))
        rental = cursor.fetchone()
        conn.close()
        
        if not rental:
            return jsonify({'success': False, 'error': 'Location non trouvée'}), 404
        
        rental = dict(rental)
        fmt = request.args.get('format', 'pdf').lower()
        
        if fmt == 'docx':
            if not DOCX_AVAILABLE:
                return jsonify({'success': False, 'error': 'python-docx non disponible'}), 500
            # Charger le modèle DOCX
            template_path = os.path.join('.', 'Modèle pour caution Location.docx')
        else:
            # Par défaut : PDF (modèle caution avec infos locataire)
            if PDF_AVAILABLE:
                pdf_buffer = generate_rental_caution_pdf(rental)
                if pdf_buffer:
                    safe_name = re.sub(r'[^\w\s-]', '', (rental.get('renter_name') or 'inconnu')).strip().replace(' ', '_')
                    filename = f'caution_location_{rental_id}_{safe_name}.pdf'
                    return send_file(
                        pdf_buffer,
                        mimetype='application/pdf',
                        as_attachment=True,
                        download_name=filename
                    )
            if not DOCX_AVAILABLE:
                return jsonify({'success': False, 'error': 'reportlab ou python-docx requis pour générer le document'}), 500
            template_path = os.path.join('.', 'Modèle pour caution Location.docx')
        
        if fmt == 'docx' or not PDF_AVAILABLE:
            template_path = os.path.join('.', 'Modèle pour caution Location.docx')
            if not os.path.exists(template_path):
                return jsonify({'success': False, 'error': 'Modèle DOCX non trouvé'}), 404
        
        doc = Document(template_path)
        
        # Formater les dates
        def format_date(date_str):
            if not date_str:
                return ''
            try:
                date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                return date_obj.strftime('%d/%m/%Y')
            except:
                return date_str
        
        # Extraire prénom et nom (format: "Prénom Nom")
        renter_name = rental['renter_name'] or ''
        name_parts = renter_name.split(' ', 1)
        prenom = name_parts[0] if len(name_parts) > 0 else ''
        nom = name_parts[1] if len(name_parts) > 1 else prenom  # Si un seul mot, c'est probablement le nom
        
        # Parser les items loués
        items_data = []
        try:
            items_data = json.loads(rental['items_data']) if rental['items_data'] else []
        except:
            items_data = []
        
        # Formater la liste des items
        items_list_str = ', '.join([
            f"{item.get('name', 'Item')} ({item.get('brand', '')} {item.get('model', '')})"
            for item in items_data
        ])
        
        # Dictionnaire de tous les placeholders possibles
        all_placeholders = {
            # Variations pour le nom
            '{{nom}}': nom,
            '{{NOM}}': nom.upper(),
            '{{prenom}}': prenom,
            '{{PRENOM}}': prenom.upper(),
            '{{nom_complet}}': renter_name,
            '{{NOM_COMPLET}}': renter_name.upper(),
            'nom_locataire': nom,
            'prenom_locataire': prenom,
            'NOM_LOCATAIRE': nom.upper(),
            'PRENOM_LOCATAIRE': prenom.upper(),
            '[NOM]': nom,
            '[PRENOM]': prenom,
            '[nom]': nom,
            '[prenom]': prenom,
            
            # Adresse
            '{{adresse}}': rental['renter_address'] or '',
            'adresse_locataire': rental['renter_address'] or '',
            '[ADRESSE]': rental['renter_address'] or '',
            '[adresse]': rental['renter_address'] or '',
            
            # Email
            '{{email}}': rental['renter_email'] or '',
            'email_locataire': rental['renter_email'] or '',
            '[EMAIL]': rental['renter_email'] or '',
            '[email]': rental['renter_email'] or '',
            
            # Téléphone
            '{{telephone}}': rental['renter_phone'] or '',
            '{{tel}}': rental['renter_phone'] or '',
            'telephone_locataire': rental['renter_phone'] or '',
            '[TELEPHONE]': rental['renter_phone'] or '',
            '[telephone]': rental['renter_phone'] or '',
            '[TEL]': rental['renter_phone'] or '',
            
            # Montants
            '{{caution}}': f"{rental['rental_deposit']:.2f} €" if rental['rental_deposit'] else '0.00 €',
            '{{montant_caution}}': f"{rental['rental_deposit']:.2f} €" if rental['rental_deposit'] else '0.00 €',
            'montant_caution': f"{rental['rental_deposit']:.2f} €" if rental['rental_deposit'] else '0.00 €',
            '[CAUTION]': f"{rental['rental_deposit']:.2f} €" if rental['rental_deposit'] else '0.00 €',
            '{{prix}}': f"{rental['rental_price']:.2f} €" if rental['rental_price'] else '0.00 €',
            '[PRIX]': f"{rental['rental_price']:.2f} €" if rental['rental_price'] else '0.00 €',
            
            # Dates
            '{{date_debut}}': format_date(rental['start_date']),
            '{{date_fin}}': format_date(rental['end_date']),
            'date_debut': format_date(rental['start_date']),
            'date_fin': format_date(rental['end_date']),
            '[DATE_DEBUT]': format_date(rental['start_date']),
            '[DATE_FIN]': format_date(rental['end_date']),
            '{{duree}}': f"{rental['rental_duration']} jour(s)" if rental['rental_duration'] else '',
            
            # Items
            '{{items}}': items_list_str,
            '{{materiel}}': items_list_str,
            '[ITEMS]': items_list_str,
            '[MATERIEL]': items_list_str,
            
            # Date du jour
            '{{date_jour}}': datetime.now().strftime('%d/%m/%Y'),
            '[DATE]': datetime.now().strftime('%d/%m/%Y'),
        }
        
        def replace_in_text(text):
            """Remplacer tous les placeholders dans un texte"""
            result = text
            
            # Remplacer les placeholders connus
            for placeholder, value in all_placeholders.items():
                result = result.replace(placeholder, str(value))
            
            # Patterns regex pour les placeholders avec "..."
            regex_patterns = [
                (r'Nom\s*[:\s]+\.{2,}', f"Nom : {renter_name}"),
                (r'Prénom\s*[:\s]+\.{2,}', f"Prénom : {prenom}"),
                (r'Adresse\s*[:\s]+\.{2,}', f"Adresse : {rental['renter_address'] or ''}"),
                (r'Email\s*[:\s]+\.{2,}', f"Email : {rental['renter_email'] or ''}"),
                (r'Téléphone\s*[:\s]+\.{2,}', f"Téléphone : {rental['renter_phone'] or ''}"),
                (r'Caution\s*[:\s]+\.{2,}', f"Caution : {rental['rental_deposit']:.2f} €" if rental['rental_deposit'] else 'Caution : 0.00 €'),
                (r'Montant\s*[:\s]+\.{2,}', f"Montant : {rental['rental_deposit']:.2f} €" if rental['rental_deposit'] else 'Montant : 0.00 €'),
                (r'Date\s+de\s+début\s*[:\s]+\.{2,}', f"Date de début : {format_date(rental['start_date'])}"),
                (r'Date\s+de\s+fin\s*[:\s]+\.{2,}', f"Date de fin : {format_date(rental['end_date'])}"),
                (r'Durée\s*[:\s]+\.{2,}', f"Durée : {rental['rental_duration']} jour(s)" if rental['rental_duration'] else 'Durée : '),
                (r'Matériel\s*[:\s]+\.{2,}', f"Matériel : {items_list_str}"),
            ]
            
            for pattern, replacement in regex_patterns:
                result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
            
            return result
        
        def replace_in_paragraph(paragraph):
            """Remplacer les placeholders dans un paragraphe en préservant la mise en forme"""
            for run in paragraph.runs:
                original_text = run.text
                new_text = replace_in_text(original_text)
                if new_text != original_text:
                    run.text = new_text
            
            # Si pas de runs ou le texte du paragraphe n'a pas été traité
            full_text = paragraph.text
            new_full_text = replace_in_text(full_text)
            if new_full_text != full_text and not paragraph.runs:
                paragraph.clear()
                paragraph.add_run(new_full_text)
        
        # Parcourir tous les paragraphes
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)
        
        # Parcourir les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)
        
        # Parcourir les en-têtes et pieds de page
        for section in doc.sections:
            # En-tête
            header = section.header
            for paragraph in header.paragraphs:
                replace_in_paragraph(paragraph)
            # Pied de page
            footer = section.footer
            for paragraph in footer.paragraphs:
                replace_in_paragraph(paragraph)
        
        # Sauvegarder dans un BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Nom de fichier sécurisé
        safe_name = re.sub(r'[^\w\s-]', '', rental['renter_name'] or 'inconnu').strip().replace(' ', '_')
        filename = f'caution_location_{rental_id}_{safe_name}.docx'
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        print(f'[API] ERREUR GET /api/rentals/{rental_id}/caution-doc: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/rental-statuses', methods=['GET'])
def get_rental_statuses():
    """Récupérer tous les statuts de location"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM rental_statuses ORDER BY name')
        rows = cursor.fetchall()
        statuses = [{
            'id': row['id'],
            'name': row['name'],
            'color': row['color']
        } for row in rows]
        conn.close()
        return jsonify({'success': True, 'statuses': statuses}), 200
    except Exception as e:
        print(f'[API] ERREUR GET /api/rental-statuses: {str(e)}')
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/rental-statuses', methods=['POST'])
def create_rental_status():
    """Créer un nouveau statut de location"""
    try:
        data = request.get_json()
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO rental_statuses (name, color, created_at)
            VALUES (?, ?, ?)
        ''', (data['name'], data.get('color', '#666'), datetime.now().isoformat()))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': cursor.lastrowid}), 201
    except Exception as e:
        print(f'[API] ERREUR POST /api/rental-statuses: {str(e)}')
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

# ==================== OCR (TESSERACT) ====================

@app.route('/api/ocr', methods=['POST'])
def ocr_image():
    """Extraire le texte d'une image avec Tesseract OCR"""
    if not OCR_AVAILABLE:
        return jsonify({
            'success': False,
            'error': 'OCR non disponible - Installez pytesseract et Tesseract'
        }), 500
    
    try:
        data = request.get_json()
        image_data = data.get('image')
        
        if not image_data:
            return jsonify({'success': False, 'error': 'Image manquante'}), 400
        
        # Décoder l'image base64
        if ',' in image_data:
            image_data = image_data.split(',')[1]
        
        image_bytes = base64.b64decode(image_data)
        image = Image.open(BytesIO(image_bytes))
        
        # Convertir en RGB si nécessaire (pour les PNG avec transparence)
        if image.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', image.size, (255, 255, 255))
            if image.mode == 'P':
                image = image.convert('RGBA')
            background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
            image = background
        elif image.mode != 'RGB':
            image = image.convert('RGB')
        
        print(f'[OCR] Traitement image {image.size}...')
        
        # Configuration Tesseract pour de meilleures performances sur les étiquettes
        # PSM 6 = Assume a single uniform block of text
        # PSM 3 = Fully automatic page segmentation (default)
        custom_config = r'--oem 3 --psm 6 -l fra+eng'
        
        # Extraire le texte
        raw_text = pytesseract.image_to_string(image, config=custom_config)
        print(f'[OCR] Texte brut extrait:\n{raw_text[:500]}...')
        
        # Parser le texte pour extraire les informations
        parsed_data = parse_ocr_text(raw_text)
        
        return jsonify({
            'success': True,
            'rawText': raw_text,
            'parsed': parsed_data
        }), 200
        
    except Exception as e:
        print(f'[OCR] Erreur: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/analyze-label-ai', methods=['POST'])
def analyze_label_ai():
    """Analyser une étiquette avec IA (OpenRouter) pour extraction intelligente des champs"""
    try:
        data = request.get_json()
        image_data = data.get('image')
        
        if not image_data:
            return jsonify({'success': False, 'error': 'Image manquante'}), 400
        
        _log('INFO', '[AI-Label] Début analyse image par IA...')
        
        # Clé API OpenRouter (à mettre en variable d'environnement en production)
        openrouter_api_key = os.environ.get('OPENROUTER_API_KEY', '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')
        
        if not openrouter_api_key or openrouter_api_key == 'YOUR_API_KEY_HERE':
            return jsonify({
                'success': False,
                'error': 'Clé API OpenRouter non configurée'
            }), 500
        
        # Récupérer les champs personnalisés
        custom_fields = []
        custom_fields_prompt = ""
        try:
            conn = get_db()
            cursor = conn.cursor()
            cursor.execute('SELECT name, field_key, field_type, options FROM custom_fields ORDER BY display_order ASC')
            rows = cursor.fetchall()
            conn.close()
            
            if rows:
                custom_fields = [{
                    'name': row['name'],
                    'fieldKey': row['field_key'],
                    'fieldType': row['field_type'],
                    'options': json.loads(row['options']) if row['options'] else None
                } for row in rows]
                
                # Ajouter les champs personnalisés au prompt
                custom_fields_json = ',\n  '.join([f'"{field["fieldKey"]}": "valeur du champ {field["name"]}"' for field in custom_fields])
                if custom_fields_json:
                    custom_fields_prompt = f",\n  {custom_fields_json}"
                
                _log('INFO', f'[AI-Label] {len(custom_fields)} champs personnalisés trouvés')
        except Exception as e:
            _log('WARN', f'[AI-Label] Impossible de récupérer les champs personnalisés: {str(e)}')
        
        # Préparer le prompt pour l'IA
        prompt = f"""Analyse cette image d'étiquette de produit et extrais les informations suivantes au format JSON strict.

Retourne UNIQUEMENT un objet JSON valide avec ces champs (mets null si information absente):
{{
  "name": "nom du produit",
  "serialNumber": "numéro de série",
  "brand": "marque",
  "model": "modèle/référence",
  "barcode": "code-barres/UPC/EAN (chiffres uniquement)",
  "description": "description courte du produit",
  "category": "catégorie (materiel, accessoire, consommable, piece_detachee, ou autre)",
  "quantity": 1{custom_fields_prompt}
}}

IMPORTANT:
- Retourne UNIQUEMENT le JSON, sans texte avant ou après
- Si une information n'est pas visible, mets null
- Pour le code-barres, extrais seulement les chiffres
- Pour la catégorie, choisis parmi: materiel, accessoire, consommable, piece_detachee, autre
- Pour les champs personnalisés, devine la valeur la plus pertinente basée sur l'image
- Sois précis et concis"""

        # Préparer la requête pour OpenRouter
        headers = {
            'Authorization': f'Bearer {openrouter_api_key}',
            'Content-Type': 'application/json',
            'HTTP-Referer': 'http://localhost:3000',  # Requis par OpenRouter
            'X-Title': 'Code Bar CRM'  # Optionnel
        }
        
        # Utiliser un modèle de vision léger et rapide
        # Modèles disponibles avec vision sur OpenRouter:
        # - openai/gpt-4o-mini (rapide, économique, fiable)
        # - anthropic/claude-3-haiku (très bon pour extraction)
        # - openai/gpt-4o (le plus puissant)
        payload = {
            'model': 'openai/gpt-4o-mini',  # Modèle OpenAI léger et fiable
            'messages': [
                {
                    'role': 'user',
                    'content': [
                        {
                            'type': 'text',
                            'text': prompt
                        },
                        {
                            'type': 'image_url',
                            'image_url': {
                                'url': image_data if image_data.startswith('data:') else f'data:image/jpeg;base64,{image_data}'
                            }
                        }
                    ]
                }
            ],
            'temperature': 0.1,  # Bas pour plus de précision
            'max_tokens': 500
        }
        
        _log('INFO', f'[AI-Label] Envoi requête à OpenRouter (modèle: {payload["model"]})...')
        
        # Appeler l'API OpenRouter
        response = requests.post(
            'https://openrouter.ai/api/v1/chat/completions',
            headers=headers,
            json=payload,
            timeout=30
        )
        
        if response.status_code != 200:
            _log('ERREUR', f'[AI-Label] Erreur API OpenRouter: {response.status_code} - {response.text}')
            return jsonify({
                'success': False,
                'error': f'Erreur API OpenRouter: {response.status_code}'
            }), 500
        
        result = response.json()
        _log('INFO', f'[AI-Label] Réponse reçue de OpenRouter')
        
        # Extraire le contenu de la réponse
        if 'choices' not in result or len(result['choices']) == 0:
            _log('ERREUR', '[AI-Label] Réponse OpenRouter invalide')
            return jsonify({
                'success': False,
                'error': 'Réponse API invalide'
            }), 500
        
        content = result['choices'][0]['message']['content']
        _log('INFO', f'[AI-Label] Contenu brut: {content[:200]}...')
        
        # Parser le JSON de la réponse
        # Parfois l'IA ajoute des balises markdown, on les enlève
        content = content.strip()
        if content.startswith('```json'):
            content = content[7:]
        if content.startswith('```'):
            content = content[3:]
        if content.endswith('```'):
            content = content[:-3]
        content = content.strip()
        
        try:
            parsed_data = json.loads(content)
            _log('INFO', f'[AI-Label] Données extraites: {json.dumps(parsed_data, ensure_ascii=False)}')
            
            # Si pas de code-barres mais un nom de produit, chercher l'UPC en ligne
            if not parsed_data.get('barcode') and parsed_data.get('name'):
                _log('INFO', f'[AI-Label] Pas de code-barres trouvé, recherche UPC en ligne pour: {parsed_data["name"]}')
                try:
                    # Construire une requête de recherche
                    search_query = parsed_data['name']
                    if parsed_data.get('brand'):
                        search_query = f"{parsed_data['brand']} {search_query}"
                    if parsed_data.get('model'):
                        search_query = f"{search_query} {parsed_data['model']}"
                    
                    _log('INFO', f'[AI-Label] Recherche UPC pour: {search_query}')
                    
                    # Essayer UPC Item DB
                    upc_api_key = os.environ.get('UPC_ITEM_DB_KEY', 'user_key$73gvintage73')
                    upc_url = f'https://api.upcitemdb.com/prod/trial/search?s={urllib.parse.quote(search_query)}'
                    upc_headers = {
                        'user_key': upc_api_key,
                        'key_type': '3scale'
                    }
                    
                    upc_response = requests.get(upc_url, headers=upc_headers, timeout=5)
                    
                    if upc_response.status_code == 200:
                        upc_data = upc_response.json()
                        if upc_data.get('items') and len(upc_data['items']) > 0:
                            first_item = upc_data['items'][0]
                            found_upc = first_item.get('upc') or first_item.get('ean')
                            if found_upc:
                                parsed_data['barcode'] = found_upc
                                _log('INFO', f'[AI-Label] UPC trouvé: {found_upc}')
                            else:
                                _log('INFO', '[AI-Label] Aucun UPC dans le premier résultat')
                        else:
                            _log('INFO', '[AI-Label] Aucun résultat UPC trouvé')
                    else:
                        _log('WARN', f'[AI-Label] Erreur recherche UPC: {upc_response.status_code}')
                    
                except Exception as search_error:
                    _log('WARN', f'[AI-Label] Erreur lors de la recherche UPC: {str(search_error)}')
                    # Continue sans UPC, ce n'est pas bloquant
            
            return jsonify({
                'success': True,
                'parsed': parsed_data,
                'rawResponse': content,
                'model': payload['model'],
                'customFields': custom_fields
            }), 200
            
        except json.JSONDecodeError as e:
            _log('ERREUR', f'[AI-Label] Erreur parsing JSON: {str(e)}')
            _log('ERREUR', f'[AI-Label] Contenu reçu: {content}')
            return jsonify({
                'success': False,
                'error': 'Réponse IA non parsable',
                'rawResponse': content
            }), 500
        
    except requests.exceptions.Timeout:
        _log('ERREUR', '[AI-Label] Timeout API OpenRouter')
        return jsonify({
            'success': False,
            'error': 'Timeout - L\'API a mis trop de temps à répondre'
        }), 500
        
    except requests.exceptions.RequestException as e:
        _log('ERREUR', f'[AI-Label] Erreur requête: {str(e)}')
        return jsonify({
            'success': False,
            'error': f'Erreur de connexion: {sanitize_error(e)}'
        }), 500
        
    except Exception as e:
        _log('ERREUR', f'[AI-Label] Erreur inattendue: {str(e)}', e)
        import traceback
        safe_traceback()
        return jsonify({
            'success': False,
            'error': sanitize_error(e)
        }), 500

def parse_ocr_text(text):
    """Parser le texte OCR pour extraire les informations utiles"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    
    result = {
        'name': None,
        'serialNumber': None,
        'brand': None,
        'model': None,
        'barcode': None,
        'description': None,
    }
    
    # Patterns pour détecter les champs
    patterns = {
        # Numéro de série
        'serialNumber': [
            r'(?:S/N|SN|Serial|N°\s*série|Numéro\s*de\s*série)[:\s]*([A-Z0-9\-]+)',
            r'(?:^|\s)([A-Z]{2,4}[0-9]{6,12})(?:\s|$)',  # Format type XX123456789
        ],
        # Code-barres / UPC / EAN
        'barcode': [
            r'(?:UPC|EAN|GTIN|Code[- ]?barre)[:\s]*(\d{8,14})',
            r'(?:^|\s)(\d{12,14})(?:\s|$)',  # 12-14 chiffres seuls
        ],
        # Modèle
        'model': [
            r'(?:Model|Modèle|Mod\.?|Ref\.?|Référence)[:\s]*([A-Z0-9\-\/\s]+)',
            r'(?:Part\s*(?:No|Number|#)|P/N)[:\s]*([A-Z0-9\-]+)',
        ],
        # Marque (liste de marques connues)
        'brand': [
            r'\b(Apple|Samsung|Sony|LG|Dell|HP|Lenovo|Asus|Acer|Microsoft|Google|DJI|GoPro|Canon|Nikon|Panasonic|Logitech|Bose|JBL|Meta|Oculus|HTC|Valve|Razer|Corsair|SteelSeries|Rode|Shure|Sennheiser|Audio-Technica|Blackmagic|Elgato|AVerMedia|OBS|Streamlabs|Insta360|Zhiyun|DJI)\b',
        ],
    }
    
    full_text = ' '.join(lines)
    
    for field, regex_list in patterns.items():
        for pattern in regex_list:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                result[field] = match.group(1).strip()
                break
    
    # Si pas de nom détecté, utiliser la première ligne significative
    if not result['name']:
        for line in lines:
            # Ignorer les lignes trop courtes ou qui ressemblent à des numéros
            if len(line) > 3 and not re.match(r'^[\d\-\/\.\s]+$', line):
                # Ignorer si c'est déjà identifié comme autre chose
                if line not in [result['serialNumber'], result['model'], result['barcode']]:
                    result['name'] = line[:100]  # Limiter la longueur
                    break
    
    # Si pas de description, utiliser les lignes non utilisées
    used_values = [v for v in result.values() if v]
    description_lines = [l for l in lines if l not in used_values and len(l) > 5]
    if description_lines:
        result['description'] = ' | '.join(description_lines[:3])  # Max 3 lignes
    
    print(f'[OCR] Données parsées: {result}')
    return result

@app.route('/api/ocr/status', methods=['GET'])
def ocr_status():
    """Vérifier si l'OCR est disponible"""
    return jsonify({
        'success': True,
        'available': OCR_AVAILABLE,
        'tesseract_path': pytesseract.pytesseract.tesseract_cmd if OCR_AVAILABLE else None
    }), 200

# ==================== COMMANDE VOCALE (IA) ====================

# Initialiser Whisper local (faster-whisper)
WHISPER_AVAILABLE = False
whisper_model = None

try:
    from faster_whisper import WhisperModel
    # Utiliser le modèle "base" pour un bon équilibre vitesse/qualité
    # Options: tiny, base, small, medium, large-v3
    print('[WHISPER] Chargement du modèle Whisper local (base)...')
    whisper_model = WhisperModel("base", device="cpu", compute_type="int8")
    WHISPER_AVAILABLE = True
    print('[WHISPER] Modèle Whisper local chargé avec succès')
except ImportError:
    print('[WHISPER] faster-whisper non installé - pip install faster-whisper')
except Exception as e:
    print(f'[WHISPER] Erreur chargement modèle: {e}')

# Vérifier si OpenAI est disponible (pour l'analyse GPT)

try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
    # Initialiser le client OpenAI (la clé sera lue depuis .env ou variable d'environnement)
    openai_api_key = os.environ.get('OPENAI_API_KEY')
    if openai_api_key:
        openai_client = OpenAI(api_key=openai_api_key)
        print('[OPENAI] Client OpenAI initialisé (pour analyse GPT)')
    else:
        OPENAI_AVAILABLE = False
        print('[OPENAI] OPENAI_API_KEY non définie - analyse GPT désactivée')
except ImportError:
    OPENAI_AVAILABLE = False
    print('[OPENAI] openai non installé - pip install openai')

@app.route('/api/voice/transcribe', methods=['POST'])
def transcribe_audio():
    """Transcrire un audio en texte avec Whisper local (faster-whisper)"""
    if not WHISPER_AVAILABLE:
        return jsonify({
            'success': False,
            'error': 'Whisper non disponible. Installez faster-whisper: pip install faster-whisper'
        }), 503
    
    try:
        # Récupérer le fichier audio
        if 'audio' not in request.files:
            return jsonify({'success': False, 'error': 'Aucun fichier audio'}), 400
        
        audio_file = request.files['audio']
        if audio_file.filename == '':
            return jsonify({'success': False, 'error': 'Fichier audio vide'}), 400
        
        print(f'[WHISPER] Transcription audio: {audio_file.filename}')
        
        # Sauvegarder temporairement le fichier
        temp_path = os.path.join('data', f'temp_audio_{datetime.now().timestamp()}.webm')
        os.makedirs('data', exist_ok=True)
        audio_file.save(temp_path)
        
        try:
            # Transcrire avec Whisper local
            segments, info = whisper_model.transcribe(
                temp_path,
                language="fr",
                beam_size=5
            )
            
            # Combiner tous les segments
            text = " ".join([segment.text for segment in segments]).strip()
            
            if not text:
                return jsonify({
                    'success': False,
                    'error': 'Aucun texte détecté dans l\'audio'
                }), 400
            
            print(f'[WHISPER] Transcription réussie: {text[:100]}...')
            print(f'[WHISPER] Langue détectée: {info.language}, probabilité: {info.language_probability:.2f}')
            
            return jsonify({
                'success': True,
                'text': text,
                'confidence': info.language_probability,
                'language': info.language
            }), 200
            
        finally:
            # Supprimer le fichier temporaire
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    except Exception as e:
        print(f'[WHISPER] Erreur transcription: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/voice/status', methods=['GET'])
def voice_status():
    """Vérifier le statut des services vocaux (Whisper + OpenAI)"""
    return jsonify({
        'whisper': {
            'available': WHISPER_AVAILABLE,
            'model': 'base (local)' if WHISPER_AVAILABLE else None
        },
        'openai': {
            'available': OPENAI_AVAILABLE,
            'note': 'Requis pour l\'analyse GPT du texte'
        }
    }), 200

@app.route('/api/voice/analyze', methods=['POST'])
def analyze_voice_command():
    """Analyser un texte avec GPT pour extraire les informations de location"""
    if not OPENAI_AVAILABLE:
        return jsonify({
            'success': False,
            'error': 'OpenAI non disponible'
        }), 503
    
    try:
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({'success': False, 'error': 'Texte vide'}), 400
        
        print(f'[VOICE] Analyse du texte: {text[:100]}...')
        
        # Créer un prompt pour GPT
        prompt = f"""Tu es un assistant qui analyse des commandes vocales pour créer des locations d'équipement.

Analyse le texte suivant et extrait les informations pour créer une location:

Texte: "{text}"

Tu dois extraire et retourner un JSON avec les champs suivants:
- items: liste d'objets avec itemId et quantity
- renterName: nom du locataire (si mentionné)
- startDate: date de début au format YYYY-MM-DD
- endDate: date de fin au format YYYY-MM-DD
- rentalPrice: prix de location (si mentionné, sinon 0)
- rentalDeposit: caution (si mentionnée, sinon 0)
- notes: notes additionnelles

IMPORTANT pour les items:
- Les identifiants d'items peuvent être: "c15", "C15", "15", "07", "A53", "AST000210", etc.
- Mets l'identifiant EXACTEMENT comme il est dit dans le champ "itemId"
- Exemples: si l'utilisateur dit "c15" -> itemId: "c15", si "07" -> itemId: "07"
- Si des quantités sont mentionnées, utilise-les, sinon mets 1
- Si le nom d'un item est mentionné au lieu d'un ID, mets-le dans le champ "name"

Pour les dates:
- Si seul le jour est mentionné, utilise le mois et l'année actuels
- "aujourd'hui" = date du jour, "demain" = jour +1, "après-demain" = jour +2

Retourne UNIQUEMENT le JSON valide, sans texte avant ou après.

Date actuelle: {datetime.now().strftime('%Y-%m-%d')}
"""
        
        # Appeler GPT (gpt-3.5-turbo = 20x moins cher que gpt-4)
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Tu es un assistant qui extrait des informations structurées depuis du texte."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3  # Basse température pour plus de précision
        )
        
        # Extraire la réponse
        result_text = response.choices[0].message.content.strip()
        print(f'[VOICE] Réponse GPT: {result_text}')
        
        # Parser le JSON
        # Nettoyer si GPT a ajouté des backticks markdown
        if result_text.startswith('```'):
            result_text = result_text.split('```')[1]
            if result_text.startswith('json'):
                result_text = result_text[4:]
            result_text = result_text.strip()
        
        result = json.loads(result_text)
        
        # Valider et nettoyer les données
        if not isinstance(result.get('items'), list) or len(result['items']) == 0:
            return jsonify({
                'success': False,
                'error': 'Aucun item détecté dans la commande vocale'
            }), 400
        
        # S'assurer que chaque item a une quantité et un identifiant
        for item in result['items']:
            if 'quantity' not in item:
                item['quantity'] = 1
            # Normaliser l'identifiant
            if 'itemId' not in item and 'serialNumber' in item:
                item['itemId'] = item['serialNumber']
            # Log pour debug
            print(f'[VOICE] Item extrait: {item}')
        
        print(f'[VOICE] Analyse réussie: {len(result["items"])} items détectés')
        
        return jsonify(result), 200
    
    except json.JSONDecodeError as e:
        print(f'[VOICE] Erreur parsing JSON: {e}')
        return jsonify({
            'success': False,
            'error': 'Impossible de parser la réponse de l\'IA'
        }), 500
    except Exception as e:
        print(f'[VOICE] Erreur analyse: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

@app.route('/api/voice/create-rental', methods=['POST'])
def create_rental_from_voice():
    """Créer une location depuis une analyse IA"""
    try:
        data = request.get_json()
        
        print(f'[VOICE] Création location depuis IA: {data}')
        
        # Récupérer les items depuis la base de données
        conn = get_db()
        cursor = conn.cursor()
        
        items_data = []
        missing_items = []
        
        for item_info in data.get('items', []):
            serial_number = item_info.get('serialNumber') or item_info.get('itemId')
            quantity = item_info.get('quantity', 1)
            
            if serial_number:
                # Chercher l'item dans la base (par serial_number, item_id ou hex_id)
                cursor.execute('''
                    SELECT serial_number, name, brand, model, item_type, quantity as stock_quantity
                    FROM items
                    WHERE serial_number = ? OR item_id = ? OR hex_id = ?
                ''', (serial_number, serial_number, serial_number.upper()))
                
                row = cursor.fetchone()
                
                if row:
                    items_data.append({
                        'serialNumber': row['serial_number'],
                        'name': row['name'],
                        'brand': row['brand'],
                        'model': row['model'],
                        'itemType': row['item_type'],
                        'quantity': quantity
                    })
                else:
                    missing_items.append(serial_number)
            else:
                # Item sans ID, on utilise juste le nom
                items_data.append({
                    'serialNumber': f'TEMP-{datetime.now().timestamp()}',
                    'name': item_info.get('name', 'Item sans nom'),
                    'quantity': quantity
                })
        
        if len(items_data) == 0:
            return jsonify({
                'success': False,
                'error': f'Aucun item trouvé. Items manquants: {", ".join(missing_items)}'
            }), 404
        
        # Valider et formater les dates (valeurs par défaut si manquantes)
        today = datetime.now()
        default_start = today.strftime('%Y-%m-%d')
        default_end = (today + timedelta(days=7)).strftime('%Y-%m-%d')
        
        start_date_str = data.get('startDate') or default_start
        end_date_str = data.get('endDate') or default_end
        
        # Nettoyer les dates (enlever les heures si présentes)
        if start_date_str and 'T' in start_date_str:
            start_date_str = start_date_str.split('T')[0]
        if end_date_str and 'T' in end_date_str:
            end_date_str = end_date_str.split('T')[0]
        
        print(f'[VOICE] Dates: {start_date_str} -> {end_date_str}')
        
        # Créer la location
        rental_data = {
            'renterName': data.get('renterName') or 'Locataire (commande vocale)',
            'renterEmail': data.get('renterEmail') or '',
            'renterPhone': data.get('renterPhone') or '',
            'renterAddress': data.get('renterAddress') or '',
            'rentalPrice': float(data.get('rentalPrice') or 0),
            'rentalDeposit': float(data.get('rentalDeposit') or 0),
            'startDate': start_date_str,
            'endDate': end_date_str,
            'itemsData': items_data,
            'notes': data.get('notes') or 'Créé depuis commande vocale'
        }
        
        # Calculer la durée
        try:
            start = datetime.strptime(rental_data['startDate'], '%Y-%m-%d')
            end = datetime.strptime(rental_data['endDate'], '%Y-%m-%d')
            rental_data['rentalDuration'] = max(1, (end - start).days)
        except Exception as date_err:
            print(f'[VOICE] Erreur parsing dates: {date_err}')
            rental_data['rentalDuration'] = 7
        
        # Déterminer le statut
        try:
            start_date = datetime.strptime(rental_data['startDate'], '%Y-%m-%d')
            rental_data['status'] = 'a_venir' if start_date > today else 'en_cours'
        except:
            rental_data['status'] = 'en_cours'
        
        # Insérer dans la base
        now = datetime.now().isoformat()
        cursor.execute('''
            INSERT INTO rentals (
                renter_name, renter_email, renter_phone, renter_address,
                rental_price, rental_deposit, rental_duration,
                start_date, end_date, status, items_data, notes,
                created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            rental_data['renterName'],
            rental_data['renterEmail'],
            rental_data['renterPhone'],
            rental_data['renterAddress'],
            rental_data['rentalPrice'],
            rental_data['rentalDeposit'],
            rental_data['rentalDuration'],
            rental_data['startDate'],
            rental_data['endDate'],
            rental_data['status'],
            json.dumps(rental_data['itemsData']),
            rental_data['notes'],
            now,
            now
        ))
        
        rental_id = cursor.lastrowid
        
        # Mettre à jour l'inventaire
        item_status = 'location_future' if rental_data['status'] == 'a_venir' else 'en_location'
        
        for item_data in rental_data['itemsData']:
            serial_number = item_data.get('serialNumber')
            quantity = item_data.get('quantity', 1)
            
            if serial_number and not serial_number.startswith('TEMP-'):
                cursor.execute('SELECT quantity FROM items WHERE serial_number = ?', (serial_number,))
                row = cursor.fetchone()
                
                if row:
                    current_quantity = row['quantity'] or 1
                    remaining_quantity = max(0, current_quantity - quantity)
                    
                    cursor.execute('''
                        UPDATE items 
                        SET status = ?,
                            quantity = ?,
                            rental_end_date = ?,
                            current_rental_id = ?,
                            last_updated = ?
                        WHERE serial_number = ?
                    ''', (
                        item_status if remaining_quantity == 0 else 'en_stock',
                        remaining_quantity,
                        rental_data['endDate'],
                        rental_id,
                        now,
                        serial_number
                    ))
        
        conn.commit()
        conn.close()
        
        # Diffuser les événements
        broadcast_event('rentals_changed', {'action': 'created', 'id': rental_id, 'source': 'voice'})
        broadcast_event('items_changed', {'action': 'updated', 'rental_id': rental_id})
        
        print(f'[VOICE] Location créée: #{rental_id}')
        
        return jsonify({
            'success': True,
            'id': rental_id,
            'itemsFound': len(items_data),
            'itemsMissing': missing_items
        }), 201
    
    except Exception as e:
        print(f'[VOICE] Erreur création location: {str(e)}')
        import traceback
        safe_traceback()
        return jsonify({'success': False, 'error': sanitize_error(e)}), 500

# ==================== HEALTH CHECK ====================

@app.route('/api/health', methods=['GET'])
def health_check():
    """Vérifier l'état du serveur"""
    return jsonify({
        'success': True,
        'status': 'healthy',
        'mode': APP_MODE,
        'database': 'connected' if os.path.exists(DB_PATH) else 'not found',
        'ocr': 'available' if OCR_AVAILABLE else 'unavailable',
        'docx': 'available' if DOCX_AVAILABLE else 'unavailable'
    }), 200

# ==================== CATCH-ALL FRONTEND (doit être après toutes les routes API) ====================

@app.route('/<path:path>')
def serve_frontend_pages(path):
    """Servir les pages du frontend (catch-all)"""
    if not FRONTEND_AVAILABLE:
        return '', 404
    
    # Ne pas intercepter les routes API
    if path.startswith('api/'):
        return '', 404
    
    # Essayer de servir le fichier directement
    file_path = os.path.join(FRONTEND_BUILD_DIR, path)
    if os.path.exists(file_path) and os.path.isfile(file_path):
        directory = os.path.dirname(file_path)
        filename = os.path.basename(file_path)
        return send_from_directory(directory, filename)
    
    # Essayer avec .html
    html_path = os.path.join(FRONTEND_BUILD_DIR, f'{path}.html')
    if os.path.exists(html_path):
        directory = os.path.dirname(html_path)
        filename = os.path.basename(html_path)
        return send_from_directory(directory, filename)
    
    # Essayer comme dossier avec index.html
    index_path = os.path.join(FRONTEND_BUILD_DIR, path, 'index.html')
    if os.path.exists(index_path):
        return send_from_directory(os.path.join(FRONTEND_BUILD_DIR, path), 'index.html')
    
    # Fallback: retourner index.html pour le routing côté client
    return send_from_directory(FRONTEND_BUILD_DIR, 'index.html')

# ==================== DÉMARRAGE ====================

import subprocess
import sys

# Configuration du démarrage (en mode dev : rechargement auto du backend à chaque modification)
FLASK_DEBUG = os.environ.get('FLASK_DEBUG', 'true' if APP_MODE == 'development' else 'false').lower() == 'true'

def build_frontend():
    """Construire le frontend Next.js si nécessaire (depuis la racine du projet)"""
    project_root = os.path.dirname(os.path.abspath(__file__))
    
    # Vérifier si le build existe déjà
    if FRONTEND_AVAILABLE:
        print("[BUILD] Frontend déjà buildé, prêt à servir.")
        return True
    
    print("[BUILD] Frontend non buildé. Construction en cours...")
    print("[BUILD] Cela peut prendre quelques minutes...")
    
    try:
        # Installer les dépendances si nécessaire
        node_modules = os.path.join(project_root, 'node_modules')
        if not os.path.exists(node_modules):
            print("[BUILD] Installation des dépendances (npm install)...")
            result = subprocess.run(
                'npm install' if sys.platform == 'win32' else ['npm', 'install'],
                cwd=project_root,
                shell=sys.platform == 'win32',
                capture_output=True,
                text=True
            )
            if result.returncode != 0:
                print(f"[BUILD] Erreur npm install: {result.stderr}")
                return False
        
        # Build le frontend
        print("[BUILD] Construction du frontend (npm run build)...")
        result = subprocess.run(
            'npm run build' if sys.platform == 'win32' else ['npm', 'run', 'build'],
            cwd=project_root,
            shell=sys.platform == 'win32',
            capture_output=True,
            text=True
        )
        if result.returncode != 0:
            print(f"[BUILD] Erreur npm run build: {result.stderr}")
            return False
        
        print("[BUILD] Frontend buildé avec succès!")
        return True
        
    except Exception as e:
        print(f"[BUILD] Erreur lors du build: {e}")
        return False

if __name__ == '__main__':
    print("=" * 60)
    print("  CODE BAR CRM - Serveur Unifié")
    print("=" * 60)
    
    # Initialiser la base de données
    init_db()
    
    # Migrer les hex_id vers le nouveau format 3 caractères
    migrate_hex_ids()
    
    # Vérifier/construire le frontend si demandé
    auto_build = os.environ.get('AUTO_BUILD', 'false').lower() == 'true'
    if not FRONTEND_AVAILABLE and auto_build:
        build_frontend()
        # Recharger la vérification
        globals()['FRONTEND_AVAILABLE'] = check_frontend_build()
    
    print("\n" + "=" * 60)
    print("  SERVEUR UNIFIE ACTIF")
    print(f"    URL:      http://localhost:{SERVER_PORT}")
    print(f"    Mode:     {APP_MODE}")
    print(f"    Frontend: {'[OK] Disponible' if FRONTEND_AVAILABLE else '[KO] Non builde'}")
    print(f"    OCR:      {'[OK] Disponible' if OCR_AVAILABLE else '[KO] Non disponible'}")
    print(f"    DOCX:     {'[OK] Disponible' if DOCX_AVAILABLE else '[KO] Non disponible'}")
    print("=" * 60)
    
    if not FRONTEND_AVAILABLE:
        print("\n[ATTENTION] Pour activer le frontend, executez (depuis la racine du projet) :")
        print("    npm install")
        print("    npm run build")
        print("    Puis relancez: python server.py")
        print("\n    Ou lancez avec AUTO_BUILD=true :")
        print("    AUTO_BUILD=true python server.py")
    
    print("\nPour arrêter le serveur: Ctrl+C")
    print("=" * 60 + "\n")
    
    # Démarrer Flask (API + Frontend sur le même port)
    # debug=True en dev : rechargement auto quand vous modifiez server.py ou les fichiers Python
    app.run(
        host='0.0.0.0',
        port=SERVER_PORT,
        debug=FLASK_DEBUG,
        use_reloader=FLASK_DEBUG,  # Redémarrage auto du serveur à chaque modification du code Python
        threaded=True  # Permettre plusieurs connexions simultanées (SSE)
    )
